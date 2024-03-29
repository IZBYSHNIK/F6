# Copyright 2024 Degtyarev Ivan

# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at

#     http://www.apache.org/licenses/LICENSE-2.0

# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

import datetime
import json
import os
import random
import requests
import sys, platform
import traceback

import matplotlib
import matplotlib.font_manager as font_manager
from PySide6 import QtCore, QtGui, QtWidgets
from PySide6.QtCore import QRunnable, Slot, QThreadPool, QUrl
from PySide6 import QtSvgWidgets
from PySide6.QtCore import QTranslator, QLibraryInfo, Signal
from PySide6.QtWidgets import QTableWidgetItem
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
from PySide6.QtMultimedia import QAudioOutput, QMediaPlayer

from docx import Document
import doc2docx

from F6_Core.students_manager import ManagerStudents
from F6_Core.user_manager import UserManager
from F6_Core.tools import ConfigManager


import math, calendar


matplotlib.use('Qt5Agg')

VERSION = '1.2.0'

QtCore.QCoreApplication.setLibraryPaths([os.path.join('PySide6', 'qt-plugins')])
LANGUAGES = ManagerStudents.crate_eternal_iter(['english', 'china', 'russia'])

defalt_settings = {
    'CURRENT_LANGUAGE': None,
    'IS_CHANGE': False,
    'ADD_FONT_SIZE': 1,
    'NAME_FONT': 'Calibri',  
}

dirname, filename = os.path.split(os.path.abspath(__file__))

CM = ConfigManager(dirname, **defalt_settings)
try:
    CM.load_config()
except FileNotFoundError:
    CM.dump_config()

DEBAG = True
BASE_PATH = dirname
CONTENT_PATH = os.path.join(dirname, 'content')

# DOCUMENTS_PATH = dirname
# if not os.path.exists(os.path.join(DOCUMENTS_PATH, 'BD')):
#     os.makedirs(os.path.join(DOCUMENTS_PATH, "BD"))
# DOCUMENTS_PATH = os.path.join(dirname, 'BD')
# print(DOCUMENTS_PATH, 'DOC_PATH')
DOCUMENTS_PATH = os.path.expanduser("~/F6")
# DOCUMENTS_PATH = os.path.expanduser("~/000")
# DOCUMENTS_PATH = dirname


USER_MANAGER = UserManager(DOCUMENTS_PATH)
MANAGER_STUDENTS = None


is_click_license = 0


class WorkerSignals(QtCore.QObject):
    '''
    Defines the signals available from a running worker thread.

    Supported signals are:

    finished
        No data

    error
        tuple (exctype, value, traceback.format_exc() )

    result
        object data returned from processing, anything

    progress
        int indicating % progress

    '''
    finished = Signal()
    error = Signal(tuple)
    result = Signal(object)
    progress = Signal(int)


class Worker(QRunnable):
    def __init__(self, fn, *args, **kwargs):
        super(Worker, self).__init__()
        # Store constructor arguments (re-used for processing)
        self.fn = fn
        self.args = args
        self.kwargs = kwargs

        self.signals = WorkerSignals()

        self.kwargs['progress_callback'] = self.signals.progress

    @Slot()  # QtCore.Slot
    def run(self):
        '''
        Initialise the runner function with passed args, kwargs.
        '''

        # Retrieve args/kwargs here; and fire processing using them
        try:
            result = self.fn(*self.args, **self.kwargs)
        except:
            traceback.print_exc()
            exctype, value = sys.exc_info()[:2]
            self.signals.error.emit((exctype, value, traceback.format_exc()))
        else:
            self.signals.result.emit(result)  # Return the result of the processing
        finally:
            self.signals.finished.emit()


class SplashScreen(QtWidgets.QSplashScreen):
    def __init__(self):
        super().__init__()
        self.setObjectName("splash_screen")
        self.setFixedSize(379, 443)
        self.move((app.primaryScreen().size().width() - self.size().width()) // 2,
                  (app.primaryScreen().size().height() - self.size().height()) // 2 - 40)

        self.logo = QtSvgWidgets.QSvgWidget(os.path.join(CONTENT_PATH, 'media', 'logo.svg'), parent=self)
        self.logo.move(65, 25)
        self.logo.setFixedSize(256, 256)
        self.logo.setObjectName("logo")

        self.progressBar = QtWidgets.QProgressBar(self)
        self.progressBar.setGeometry(QtCore.QRect(0, 400, 381, 51))
        self.progressBar.setProperty("value", 0)
        self.progressBar.setTextVisible(False)
        self.progressBar.setObjectName("progressBar")
        self.version = QtWidgets.QLabel(self)

        self.version.setGeometry(QtCore.QRect(379 - 61 - 10, 443 - 41, 61, 41))
        font = QtGui.QFont()
        font.setItalic(True)
        self.version.setFont(font)
        self.version.setAlignment(QtCore.Qt.AlignmentFlag.AlignRight | QtCore.Qt.AlignmentFlag.AlignVCenter)
        self.version.setLayoutDirection(QtCore.Qt.LayoutDirection.RightToLeft)
        self.version.setObjectName("version")
        self.message = QtWidgets.QLabel(self)
        self.message.setGeometry(QtCore.QRect(10, 360, 351, 20))
        self.message.setObjectName("message")
        self.progressBar.raise_()
        self.logo.raise_()
        self.version.raise_()
        self.message.raise_()

        self.retranslateUi()
        QtCore.QMetaObject.connectSlotsByName(self)

    def retranslateUi(self):
        self.version.setText(str(VERSION))
        self.message.setText(self.tr("Идет загрузка программы, ожидайте ..."))


class LicenseWindow(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super(LicenseWindow, self).__init__(parent=parent)
        self.setObjectName("Form")
        self.resize(600, 700)
        self.setWindowIcon(QtGui.QIcon(os.path.join(CONTENT_PATH, 'media', 'logo.svg')))
        self.setFixedSize(QtCore.QSize(600, 700))
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.logo = QtSvgWidgets.QSvgWidget(os.path.join(CONTENT_PATH, 'media', 'logo.svg'))
        self.logo.setFixedSize(QtCore.QSize(50, 50))

        self.logo.setObjectName("label_2")
        self.verticalLayout_2.addWidget(self.logo, QtCore.Qt.AlignmentFlag.AlignVCenter,
                                        QtCore.Qt.AlignmentFlag.AlignHCenter)
        self.label = QtWidgets.QLabel(self)
        font = QtGui.QFont()
        font.setPointSize(16 + CM.ADD_FONT_SIZE)
        self.label.setFont(font)
        self.label.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.label.setObjectName("label")
        self.verticalLayout_2.addWidget(self.label)
        self.textBrowser = QtWidgets.QTextBrowser(self)
        font = QtGui.QFont()
        font.setPointSize(14 + CM.ADD_FONT_SIZE)
        self.textBrowser.setFont(font)
        self.textBrowser.setStyleSheet('border: none; background-color: rgba(0, 0, 0, 0)')
        self.textBrowser.setObjectName("textBrowser")
        self.verticalLayout_2.addWidget(self.textBrowser)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setContentsMargins(-1, 8, -1, -1)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.pushButton = QtWidgets.QPushButton(self)
        self.pushButton.setObjectName("pushButton")
        self.horizontalLayout.addWidget(self.pushButton)
        self.verticalLayout_2.addLayout(self.horizontalLayout)
        self.gw = GratutudeWindow()

        self.retranslateUi()
        QtCore.QMetaObject.connectSlotsByName(self)

        self.pushButton.clicked.connect(self.click_OK)
        self.init_license(os.path.join(BASE_PATH, 'LICENSE'))
        self.logo.mousePressEvent = self.click_logo

    def retranslateUi(self):
        self.setWindowTitle(self.tr("Лицензия"))
        self.label.setText(self.tr("Лицензионное соглашение"))
        self.pushButton.setText(self.tr("ОК"))

        self.gw.retranslateUi()

    def click_OK(self):
        global is_click_license
        is_click_license = 1
        self.close()

    def init_license(self, file_puth=None):
        with open(file_puth, 'r', encoding='Windows-1251') as f:
            self.textBrowser.setText(''.join(f.readlines()))

    def click_logo(self, e):
        self.gw.show()


class GratutudeWindow(QtWidgets.QWidget):
    def __init__(self, parent=None):
        super(GratutudeWindow, self).__init__(parent=parent)
        self.setObjectName("Form")
        self.resize(600, 700)
        self.setWindowIcon(QtGui.QIcon(os.path.join(CONTENT_PATH, 'media', 'logo.svg')))
        self.setFixedSize(QtCore.QSize(600, 700))
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.logo = QtSvgWidgets.QSvgWidget(os.path.join(CONTENT_PATH, 'media', 'heart.svg'))
        self.logo.setFixedSize(QtCore.QSize(50, 50))

        self.logo.setObjectName("label_2")
        self.verticalLayout_2.addWidget(self.logo, QtCore.Qt.AlignmentFlag.AlignVCenter,
                                        QtCore.Qt.AlignmentFlag.AlignHCenter)

        self.textBrowser = QtWidgets.QTextBrowser(self)
        font = QtGui.QFont()
        font.setPointSize(14 + CM.ADD_FONT_SIZE)
        self.textBrowser.setFont(font)
        self.textBrowser.setObjectName("textBrowser")
        self.textBrowser.setStyleSheet('border: none; background-color: rgba(0, 0, 0, 0)')
        self.verticalLayout_2.addWidget(self.textBrowser)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setContentsMargins(-1, 8, -1, -1)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.pushButton = QtWidgets.QPushButton(self)
        self.pushButton.setObjectName("pushButton")
        self.horizontalLayout.addWidget(self.pushButton)
        self.verticalLayout_2.addLayout(self.horizontalLayout)

        self.retranslateUi()
        QtCore.QMetaObject.connectSlotsByName(self)

        self.pushButton.clicked.connect(self.click_OK)
        self.logo.mousePressEvent = self.click_OK

    def retranslateUi(self):
        self.setWindowTitle(self.tr("Благодарность"))
        self.textBrowser.setHtml(self.tr(
            '<h1 style="text-align: center; color: red;">Благодарность</h1> <p style="text-align: center;">Выражаю огромную благодарность всем тем, кто постоянно окружал все это время.</p><p style="text-align: center;">Отдельная благодарность: <br>моей маме - <b>Макаровой Наталье Алексеевне</b>, <br>куратору работы – <b>Кузьминой Ирине Александровне</b>, <br>лучшему педагогу -  <b>Поповой Наталии Евгеньевне</b>,<br>моим соседям  – <b>Амангильдину Рамису и Ильясову Айсару</b>, <br>тестировщикам – <b>Кубагушеву Искандеру</b>, <br> <b>Александрову Александру</b><br> и <b>Владиславу Васильеву</b>.'))
        self.pushButton.setText(self.tr("ОК"))

    def click_OK(self, e=None):
        global is_click_license
        is_click_license = 1
        self.close()


class Regist(QtWidgets.QWidget):
    def __init__(self):
        self.status = 0
        super().__init__()
        self.setWindowIcon(QtGui.QIcon(os.path.join(CONTENT_PATH, 'media', 'logo.svg')))
        self.setObjectName("Regist")
        self.setFixedSize(442, 700)
        self.setStyleSheet("QPushButton{\n"
                           "font-size: 20px;\n"
                           "color: white;\n"
                           "}\n"
                           "\n"
                           "QPushButton:hover {\n"
                           "    border: 6px solid green;\n"
                           "    font: 21px;\n"
                           "\n"
                           "}\n"
                           "\n"
                           "Form {\n"
                           "background: #f9f8f4;\n"
                           "}\n"
                           "\n"
                           "QLabel#regist_lable{\n"
                           "    font-size: 45px;\n"
                           "font-family: Myriad Pro;\n"
                           "text-align: center;\n"
                           "background: #bcbbb7;\n"
                           "}\n"
                           "\n"
                           "QLineEdit {\n"
                           "border: none;\n"
                           "color:white;\n"
                           "padding-left: 15px;\n"
                           "padding-right: 15px;\n"
                           "font: 16px;\n"
                           "background: black;\n"
                           "}\n"
                           "")
        self.regist_lable = QtWidgets.QLabel(self)
        self.regist_lable.setEnabled(True)
        self.regist_lable.setGeometry(QtCore.QRect(0, -10, 442, 100))
        self.regist_lable.setStyleSheet("")
        self.regist_lable.setTextFormat(QtCore.Qt.TextFormat.RichText)
        self.regist_lable.setScaledContents(False)
        self.regist_lable.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.regist_lable.setObjectName("regist_lable")
        self.login_edit = QtWidgets.QLineEdit(self)
        self.login_edit.setGeometry(QtCore.QRect(41, 150, 360, 40))
        font = QtGui.QFont()
        font.setPointSize(-1)
        font.setBold(False)
        font.setItalic(False)
        self.login_edit.setFont(font)
        self.login_edit.setObjectName("login_edit")
        self.create_push = QtWidgets.QPushButton(self)
        self.create_push.setGeometry(QtCore.QRect(60, 630, 151, 61))
        self.create_push.setCursor(QtGui.QCursor(QtCore.Qt.CursorShape.ArrowCursor))
        self.create_push.setAutoFillBackground(False)
        self.create_push.setStyleSheet("background: #2ac11a;\n"
                                       "border: none;\n"
                                       "\n"
                                       "\n"
                                       "\n"
                                       "\n"
                                       "")
        self.create_push.setCheckable(False)
        self.create_push.setAutoRepeat(False)
        self.create_push.setDefault(False)
        self.create_push.setObjectName("create_push")
        self.cancel_push = QtWidgets.QPushButton(self)
        self.cancel_push.setGeometry(QtCore.QRect(220, 630, 151, 61))
        self.cancel_push.setStyleSheet("background: #2ac11a;\n"
                                       "border: none;")
        self.cancel_push.setObjectName("cancel_push")
        self.password_edit = QtWidgets.QLineEdit(self)
        self.password_edit.setGeometry(QtCore.QRect(41, 240, 360, 40))
        self.password_edit.setFocusPolicy(QtCore.Qt.FocusPolicy.StrongFocus)
        self.password_edit.setObjectName("password_edit")
        self.login_lable = QtWidgets.QLabel(self)
        self.login_lable.move(20, 100)
        font = QtGui.QFont()
        font.setPointSize(14)
        self.login_lable.setFont(font)
        self.login_lable.setObjectName("login_lable")
        self.password_lable = QtWidgets.QLabel(self)
        self.password_lable.move(20, 190)
        font = QtGui.QFont()
        font.setPointSize(14)
        self.password_lable.setFont(font)
        self.password_lable.setObjectName("password_lable")
        self.password_edit_2 = QtWidgets.QLineEdit(self)
        self.password_edit_2.setGeometry(QtCore.QRect(40, 330, 360, 40))
        self.password_edit_2.setFocusPolicy(QtCore.Qt.FocusPolicy.StrongFocus)
        self.password_edit_2.setObjectName("password_edit_2")
        self.password_lable_2 = QtWidgets.QLabel(self)
        self.password_lable_2.move(30, 280)
        font = QtGui.QFont()
        font.setPointSize(14)
        self.password_lable_2.setFont(font)
        self.password_lable_2.setObjectName("password_lable_2")
        self.teamleader_lable = QtWidgets.QLabel(self)
        self.teamleader_lable.move(240, 380)
        font = QtGui.QFont()
        font.setPointSize(11)
        self.teamleader_lable.setFont(font)
        self.teamleader_lable.setObjectName("teamleader_lable")
        self.teamleader_edit = QtWidgets.QLineEdit(self)
        self.teamleader_edit.setGeometry(QtCore.QRect(240, 430, 160, 40))
        self.teamleader_edit.setFocusPolicy(QtCore.Qt.FocusPolicy.StrongFocus)
        self.teamleader_edit.setObjectName("teamleader_edit")
        self.office_name_edit = QtWidgets.QLineEdit(self)
        self.office_name_edit.setGeometry(QtCore.QRect(40, 430, 160, 40))
        self.office_name_edit.setFocusPolicy(QtCore.Qt.FocusPolicy.StrongFocus)
        self.office_name_edit.setObjectName("office_name_edit")
        self.office_name_lable = QtWidgets.QLabel(self)
        self.office_name_lable.move(40, 380)
        font = QtGui.QFont()
        font.setPointSize(11)
        self.office_name_lable.setFont(font)
        self.office_name_lable.setObjectName("office_name_lable")
        self.group_edit = QtWidgets.QLineEdit(self)
        self.group_edit.setGeometry(QtCore.QRect(40, 520, 160, 40))
        font = QtGui.QFont()
        font.setPointSize(-1)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(QtGui.QFont.Weight(50))
        self.group_edit.setFont(font)
        self.group_edit.setObjectName("group_edit")
        self.specialization_edit = QtWidgets.QLineEdit(self)
        self.specialization_edit.setGeometry(QtCore.QRect(240, 520, 160, 40))
        font = QtGui.QFont()
        font.setPointSize(-1)
        font.setBold(False)
        font.setItalic(False)
        font.setWeight(QtGui.QFont.Weight(50))
        self.specialization_edit.setFont(font)
        self.specialization_edit.setObjectName("specialization_edit")
        self.group_lable = QtWidgets.QLabel(self)
        self.group_lable.setGeometry(QtCore.QRect(40, 470, 311, 51))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.group_lable.setFont(font)
        self.group_lable.setObjectName("group_lable")
        self.specialization_lable = QtWidgets.QLabel(self)
        self.specialization_lable.setGeometry(QtCore.QRect(240, 470, 311, 51))
        font = QtGui.QFont()
        font.setPointSize(11)
        self.specialization_lable.setFont(font)
        self.specialization_lable.setObjectName("specialization_lable")
        self.message_regist = QtWidgets.QTextBrowser(self)
        self.message_regist.setGeometry(40, 550 + 30, 351, 41)
        self.message_regist.setStyleSheet(
            'border: none; color: red; font: 14px; background-color: rgba(249, 248, 244, 0);')

        self.logo = QtSvgWidgets.QSvgWidget(self)
        self.logo.load(os.path.join(CONTENT_PATH, 'media', 'language.svg'))
        self.logo.setFixedSize(QtCore.QSize(50, 50))
        self.logo.move(362, 100)
        self.logo.mousePressEvent = set_language

        self.retranslateUi()
        self.add_functions()
        QtCore.QMetaObject.connectSlotsByName(self)

    def retranslateUi(self):

        self.setWindowTitle(self.tr("Регистрация"))
        self.regist_lable.setText(self.tr("РЕГИСТРАЦИЯ"))
        self.create_push.setText(self.tr("Создать"))
        self.cancel_push.setText(self.tr("Отмена"))
        self.login_lable.setText(self.tr("*Имя пользователя"))
        self.password_lable.setText(self.tr("*Пароль"))
        self.password_lable_2.setText(self.tr("*Повторите пароль"))
        self.teamleader_lable.setText(self.tr("ФИО Кл.руководителя"))
        self.office_name_lable.setText(self.tr("Своё ФИО"))
        self.group_lable.setText(self.tr("Группа"))
        self.specialization_lable.setText(self.tr('Специализация'))

    def add_functions(self):
        self.create_push.clicked.connect(self.click_create_push)
        self.cancel_push.clicked.connect(self.click_cancel_push)

    def check_lables(self):

        kwargs = {}
        if not USER_MANAGER.USER_CLASS.is_valid_username(self.login_edit.text()):
            raise ValueError(self.tr('Имя пользователя должно быть строкой из 3-20 букв или числовых символов'))
        if not USER_MANAGER.USER_CLASS.is_valid_password(self.password_edit.text()):
            raise ValueError(self.tr('Пароль должен быть строкой из 4-30 символов'))
        if self.login_edit.text() in USER_MANAGER.users_id.values():
            raise ValueError(self.tr('Имя пользователя занято'))
        if not self.password_edit.text() == self.password_edit_2.text():
            raise ValueError(self.tr('Не совпадают пароли'))
        if self.office_name_edit.text():
            if USER_MANAGER.USER_CLASS.is_valid_fio(self.office_name_edit.text()):
                kwargs['offical_name'] = self.office_name_edit.text()
            else:
                raise ValueError(
                    self.tr(
                        'ФИО должно состоять только из букв и быть из 3 частей, каждая из которых не менее 2 символов'))
        if self.teamleader_edit.text():
            if USER_MANAGER.USER_CLASS.is_valid_fio(self.teamleader_edit.text()):
                kwargs['teamleader'] = USER_MANAGER.USER_CLASS.check_fio(self.teamleader_edit.text())
            else:
                raise ValueError(
                    self.tr(
                        'ФИО должно состоять только из букв и быть из 3 частей, каждая из которых не менее 2 символов'))
        if self.group_edit.text():
            kwargs['group'] = self.group_edit.text()
        if self.specialization_edit.text():
            kwargs['specialization'] = self.specialization_edit.text()
        return self.login_edit.text(), self.password_edit.text(), kwargs

    def click_create_push(self):
        self.message_regist.setText('')
        try:
            username, password, kwargs = self.check_lables()
            user = USER_MANAGER.USER_CLASS(username, password, kwargs)
            USER_MANAGER.link_user_by_obj(user)
            USER_MANAGER.save_users()


        except BaseException as message:
            self.message_regist.setText('*' + str(message))
        else:
            self.password_edit.clear()
            self.password_edit_2.clear()
            self.group_edit.clear()
            self.specialization_edit.clear()
            self.login_edit.clear()
            self.teamleader_edit.clear()
            self.office_name_edit.clear()
            self.status = 1
            self.close()

    def click_cancel_push(self):
        self.status = 2
        self.close()


class Auth(QtWidgets.QWidget):

    def __init__(self):
        self.status = 0
        super(Auth, self).__init__()
        self.setWindowIcon(QtGui.QIcon(os.path.join(CONTENT_PATH, 'media', 'logo.svg')))
        self.setFixedSize(442, 580)

        self.setStyleSheet("QPushButton{font: 18px; border: none; color: white; padding: 0px; margin: 0px}\n"
                           "QPushButton:hover {font-size: 21px;}\n"
                           "Form {background: #f9f8f4;}\n"
                           "QLabel#auth{font-size: 45px;font-family: Myriad Pro;text-align: center;background: #bcbbb7;}\n"
                           "QLabel#login_lable{font-size: 35px;text-align: center;}\n"
                           "QLabel#password_lable{font-size: 35px;text-align: center;}\n"
                           "QLabel{font-size: 14px;text-align: center;}\n"
                           "QLineEdit {border: 2px solid black;color:white;padding-left: 15px;font: 16px;background: black;}\n")
        self.auth = QtWidgets.QLabel(self)
        self.auth.setEnabled(True)
        self.auth.setGeometry(QtCore.QRect(0, -10, 442, 100))
        self.auth.setMinimumSize(QtCore.QSize(442, 100))
        self.auth.setMaximumSize(QtCore.QSize(442, 70))
        self.auth.setStyleSheet("")
        self.auth.setTextFormat(QtCore.Qt.TextFormat.RichText)
        self.auth.setScaledContents(False)
        self.auth.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.auth.setObjectName("auth")

        self.logo = QtSvgWidgets.QSvgWidget(self)
        self.logo.load(os.path.join(CONTENT_PATH, 'media', 'language.svg'))
        self.logo.setFixedSize(QtCore.QSize(50, 50))
        self.logo.move(362, 100)

        self.in_push = QtWidgets.QPushButton(self)
        self.in_push.setGeometry(QtCore.QRect(70, 480, 151, 61))
        self.in_push.setCursor(QtGui.QCursor(QtCore.Qt.CursorShape.ArrowCursor))
        self.in_push.setAutoFillBackground(False)
        self.in_push.setStyleSheet("background: #2ac11a;\n")
        self.in_push.setCheckable(False)
        self.in_push.setAutoRepeat(False)
        self.in_push.setDefault(False)
        self.in_push.setObjectName("in_push")
        self.regist_push = QtWidgets.QPushButton(self)
        self.regist_push.setGeometry(QtCore.QRect(230, 480, 151, 61))
        self.regist_push.setStyleSheet("background: #2ac11a;\n")
        self.regist_push.setObjectName("regist_push")
        self.password_edit = QtWidgets.QLineEdit(self)
        self.password_edit.setGeometry(QtCore.QRect(60, 300, 351, 41))
        self.password_edit.setFocusPolicy(QtCore.Qt.FocusPolicy.StrongFocus)
        self.password_edit.setEchoMode(QtWidgets.QLineEdit.EchoMode.Password)
        self.password_edit.setObjectName("password_edit")
        self.login_lable = QtWidgets.QLabel(self)
        self.login_lable.move(70, 120)
        self.login_lable.setFixedWidth(250)
        self.login_lable.setObjectName("login_lable")
        self.password_lable = QtWidgets.QLabel(self)
        self.password_lable.move(70, 235)
        self.password_lable.setFixedWidth(300)
        self.password_lable.setObjectName("password_lable")
        self.message_auth = QtWidgets.QTextBrowser(self)
        self.message_auth.setGeometry(60, 351 + 20, 351, 41)

        self.eye = QtSvgWidgets.QSvgWidget(self)
        self.eye.load(os.path.join(CONTENT_PATH, 'media', 'eye_open.svg'))
        self.eye.setFixedSize(QtCore.QSize(40, 40))
        self.eye.move(365, 300)

        self.licensewindow = LicenseWindow()
        self.license_link = QtWidgets.QPushButton(self)
        self.license_link.setStyleSheet("color: #0A5F38;")
        self.license_link.setGeometry(QtCore.QRect(0, 540, 442, 40))

        self.message_auth.setStyleSheet(
            'border: none; color: red; font: 14px; background-color: rgba(249, 248, 244, 0);')

        self.spin_box = QtWidgets.QComboBox(self)
        self.spin_box.setGeometry(QtCore.QRect(60, 180, 351, 41))
        self.spin_box.setStyleSheet("""QComboBox {color:white;font: 16px;background: black;}\n""")

        self.retranslateUi()
        QtCore.QMetaObject.connectSlotsByName(self)

        self.update_users()

        self.in_push.clicked.connect(self.click_auth_push)
        self.regist_push.clicked.connect(self.click_regis_push)
        self.license_link.clicked.connect(self.click_license)
        self.logo.mousePressEvent = set_language
        self.eye.mousePressEvent = self.set_show_password
        self.is_open = 0

    def retranslateUi(self):
        self.setWindowTitle(self.tr("Авторизация"))
        self.auth.setText(self.tr("АВТОРИЗАЦИЯ"))
        self.in_push.setText(self.tr("Вход"))
        self.regist_push.setText(self.tr("Регистрация"))
        self.login_lable.setText(self.tr("Логин"))
        self.password_lable.setText(self.tr("Пароль"))
        self.license_link.setText('©2024 Degtyarev Ivan')

        self.licensewindow.retranslateUi()

    def set_show_password(self, e):
        if not self.is_open:
            self.password_edit.setEchoMode(QtWidgets.QLineEdit.EchoMode.Normal)
            self.is_open = 1
            self.eye.load(os.path.join(CONTENT_PATH, 'media', 'eye_close.svg'))
        else:
            self.password_edit.setEchoMode(QtWidgets.QLineEdit.EchoMode.Password)
            self.is_open = 0
            self.eye.load(os.path.join(CONTENT_PATH, 'media', 'eye_open.svg'))

    def keyPressEvent(self, e):
        k = e.key()
        super().keyPressEvent(e)
        if k == 16777220:
            self.click_auth_push()


    def update_users(self):
        self.spin_box.clear()
        list_user = list(map(lambda x: f'{x[1]} ({x[0]})', USER_MANAGER.users_id.items()))
        if USER_MANAGER.parametrs.get('LastUser'):
            username = USER_MANAGER.users_id.get(USER_MANAGER.parametrs.get('LastUser'))
            if username in list_user:
                list_user.remove(username)
                list_user.insert(0, username)
        for i in list_user:
            self.spin_box.addItem(QtGui.QIcon(os.path.join(CONTENT_PATH, 'media', 'student_icon.svg')), i)

    def checking_log_password(self, login: str, password: str) -> tuple: 
        if not USER_MANAGER.USER_CLASS.is_valid_username(login):
            raise ValueError(self.tr('Имя пользователя должено быть строкой из 3-20 букв или числовых символов'))
        if not USER_MANAGER.USER_CLASS.is_valid_password(password):
            raise ValueError(self.tr('Пароль должен быть строкой из 4-30 символов'))
        return login, password

    def click_auth_push(self):

        try:
            id_ = list(USER_MANAGER.users_id)[self.spin_box.currentIndex()]          
            login = USER_MANAGER.users_id[id_]
            username, password = self.checking_log_password(login, self.password_edit.text())
        except BaseException as message:
            self.message_auth.setText(str(message))

        else:
            try:
                USER_MANAGER.link_user_by_username(username, password, id_=id_)
            except (FileNotFoundError, json.JSONDecodeError):
                self.message_auth.setText(self.tr('*Файлы пользователя повреждены'))
            except ValueError:
                self.message_auth.setText(self.tr('*Неверный пароль'))
            except BaseException:
                self.message_auth.setText(self.tr('*Внутренняя ошибка программы'))
            else:
                USER_MANAGER.user.user_id = id_
                self.password_edit.clear()
                self.status = 1
                self.close()

    def click_regis_push(self):
        self.status = 3
        self.close()

    def click_license(self):
        self.license_link.setStyleSheet("color: #2F4538;")
        self.licensewindow.show()


class ScheduleLoadWindow(QtWidgets.QDialog):
    def __init__(self, *args, parent=None):
        super().__init__(*args)
        self.parent = parent
        self.setObjectName(u"ScheduleLoadWindow")
        self.setModal(True)
        self.resize(389, 393)
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_2.setObjectName(u"verticalLayout_2")
        self.verticalLayout = QtWidgets.QVBoxLayout()
        self.verticalLayout.setObjectName(u"verticalLayout")
        self.search_label = QtWidgets.QLabel(self)
        self.search_label.setObjectName(u"search_label")

        self.verticalLayout.addWidget(self.search_label)

        self.search_lineEdit = QtWidgets.QLineEdit(self)
        self.search_lineEdit.setObjectName(u"search_lineEdit")

        self.verticalLayout.addWidget(self.search_lineEdit)

        self.period_label = QtWidgets.QLabel(self)
        self.period_label.setObjectName(u"period_label")
        self.verticalLayout.addWidget(self.period_label)

        self.period_lineEdit = QtWidgets.QDateEdit(self)
        self.period_lineEdit.setDateTime(QtCore.QDateTime(QtCore.QDate(MANAGER_STUDENTS.period[1], MANAGER_STUDENTS.period[0], 2), QtCore.QTime(0, 0, 0)))
        self.period_lineEdit.setObjectName("dateEdit")
        self.period_lineEdit.setDisplayFormat("MM.yyyy")
        self.period_lineEdit.setObjectName(u"period_lineEdit")
        self.verticalLayout.addWidget(self.period_lineEdit)
        # self.period_lineEdit.setEnabled(False)



        self.path_label = QtWidgets.QLabel(self)
        self.path_label.setObjectName(u"path_label")
        self.verticalLayout.addWidget(self.path_label)

        self.path_lineEdit = QtWidgets.QPushButton(self)
        self.path_lineEdit.setObjectName(u"path_lineEdit")
        self.verticalLayout.addWidget(self.path_lineEdit)



        self.is_convert_checkBox = QtWidgets.QCheckBox(self)
        self.is_convert_checkBox.setObjectName(u"is_convert_lineEdit")

        self.verticalLayout.addWidget(self.is_convert_checkBox)

        self.is_delete_original_files_checkBox = QtWidgets.QCheckBox(self)
        self.is_convert_checkBox.setObjectName(u"is_delete_original_files_checkBox")
        self.is_convert_checkBox.setChecked(True)

        self.verticalLayout.addWidget(self.is_delete_original_files_checkBox)

        self.message = QtWidgets.QLabel(self)
        self.message.setObjectName(u"message")

        self.verticalLayout.addWidget(self.message)

        self.table_layout = QtWidgets.QVBoxLayout(self)
        self.table_layout.setObjectName(u"table_layout")
        self.verticalLayout.addLayout(self.table_layout)

        self.verticalSpacer = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                                    QtWidgets.QSizePolicy.Policy.Expanding)

        self.verticalLayout.addItem(self.verticalSpacer)

        self.progressBar = QtWidgets.QProgressBar(self)
        self.progressBar.setObjectName(u"progressBar")
        self.progressBar.setValue(0)
        self.progressBar.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.progressBar.setTextVisible(True)

        self.verticalLayout.addWidget(self.progressBar)

        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName(u"horizontalLayout")
        self.search_button = QtWidgets.QPushButton()
        self.search_button.setObjectName(u"search_button")
        self.horizontalLayout.addWidget(self.search_button)
        self.load_data_button = QtWidgets.QPushButton()
        self.load_data_button.setObjectName(u"load_data_button")
        self.load_data_button.setEnabled(False)

        self.horizontalLayout.addWidget(self.load_data_button)

        self.verticalLayout.addLayout(self.horizontalLayout)




        self.verticalLayout_2.addLayout(self.verticalLayout)

        self.threadpool = QThreadPool()

        # self.update_table()
        self.is_stop = False
        self.is_load = False
        self.result = {}

        self.retranslateUi()
        self.add_function()

    def retranslateUi(self):
        self.setWindowTitle(self.tr(u"Загрузить расписание"))
        self.search_label.setText(self.tr(u"Искать"))
        self.path_label.setText(self.tr(u"Путь"))
        self.period_label.setText(self.tr(u'Период'))
        # self.message.setText(self.tr(u"TextLabel"))
        self.load_data_button.setText(self.tr(u"Загрузить"))
        self.search_button.setText(self.tr(u"Поиск"))
        self.is_convert_checkBox.setText(self.tr(u"Конвертировать doc в docx"))
        self.is_delete_original_files_checkBox.setText(self.tr(u"Удалить исходные файлы doc"))
        self.path_lineEdit.setText(self.tr(u"Путь ..."))

        self.search_lineEdit.setText(USER_MANAGER.user.parametrs.get('group'))

    def add_function(self):
        self.load_data_button.clicked.connect(self.click_load_data_button)
        self.path_lineEdit.clicked.connect(self.click_load_path)
        self.search_button.clicked.connect(self.click_search_button)
        self.is_convert_checkBox.stateChanged.connect(self.convert_status)

    def click_load_path(self):
        file_name = QtWidgets.QFileDialog.getExistingDirectoryUrl(self, self.tr(
            "Выберите путь и имя файла для сохранения."))
        self.path = QtCore.QUrl(file_name.url()).toLocalFile()
        self.path_lineEdit.setText(self.path)

    def convert_status(self):
        if not self.is_convert_checkBox.isChecked():
            self.is_delete_original_files_checkBox.hide()
        else:
            self.is_delete_original_files_checkBox.show()

    def set_progressBar(self, value):
        self.progressBar.setValue(value)

    def get_result(self, s):
        self.search_button.setEnabled(True)
        if s:
            self.load_data_button.setEnabled(True)
            self.result = s
        self.update_table(s)

    def get_load_error(self, e):
        self.search_button.setEnabled(True)
        self.message.setText(self.tr('Внутренняя ошибка. Проверьте верность заполнения полей.'))

    def click_search_button(self):
        self.is_load = False
        self.search_button.setEnabled(False)
        self.load_data_button.setEnabled(False)
        self.message.clear()

        self.w = Worker(self.load_files, [self.period_lineEdit.date().month(), self.period_lineEdit.date().year()])
        self.w.signals.progress.connect(self.set_progressBar)
        self.w.signals.result.connect(self.get_result)
        self.w.signals.error.connect(self.get_load_error)

        self.threadpool.start(self.w)

    def load_files(self, period, progress_callback, ):
        search = ''.join(self.search_lineEdit.text().strip().split())
        months = {}
        if not search:
            return months
        path_folder_input, _, files = \
            list(os.walk(self.path))[0]

        doc_files = list(filter(lambda x: x.endswith('.doc'), files))

        if self.is_convert_checkBox.isChecked():
            for i in range(len(doc_files)):
                progress_callback.emit(i / len(doc_files) / 2 * 100)
                if self.is_stop:
                    return {}
                path_file = os.path.join(path_folder_input, doc_files[i])
                try:
                    doc2docx.convert(path_file)
                except BaseException as e:
                    self.message.setText(str(e))
                if self.is_delete_original_files_checkBox.isChecked():
                    os.remove(path_file)
            path_folder_input, _, files = \
                list(os.walk(self.path))[0]
        else:
            if self.is_stop:
                return {}
            progress_callback.emit(50)

        files = list(filter(lambda x: x.endswith('.docx'), files))

        for f in range(len(files)):
            if self.is_stop:
                return {}
            progress_callback.emit(f / len(files) / 2 * 100 + 50)
            try:
                doc = Document(os.path.join(path_folder_input, files[f]))
                for paragraph in doc.paragraphs:
                    if len(paragraph.text) == 0:
                        p = paragraph._element
                        p.getparent().remove(p)
                        p._p = p._element = None
                month = doc.paragraphs[1].text.split()[1]

                dd, mm, yyyy = month.split('.')
                if [int(mm), int(yyyy)] == list(period):
                    for i in range(len(doc.tables[0].rows)):
                        for j in range(len(doc.tables[0].rows[i].cells)):

                            if i <= 1:
                                continue
                            if j == 0:
                                group = ''.join(doc.tables[0].rows[i].cells[0].text.strip().split()).lower()
                                if not group.lower() == search.lower():
                                    is_skip = True
                                    continue
                                else:
                                    is_skip = False

                                # print(i, j)

                            else:
                                if not is_skip:
                                    r = doc.tables[0].rows[i].cells[j].text.split('\n')

                                    if len(r) == 1 and not r[0].startswith('---'):
                                        r = r[0].split()
                                        r1 = ''.join(r[:-3])
                                        r2 = ' '.join(r[-3:-1])
                                        r = [r1, r2]
                                    if len(r) > 2 and (r[-1].upper().startswith('АУД') or r[-1].upper().startswith('ДОТ')):
                                        del r[-1]

                                    if len(r) == 2 and (
                                            r[-1].upper().find('АУД.') != -1 or r[-1].upper().find('ДОТ') != -1):
                                        r[-1] = ' '.join(r[-1].split()[:-1])

                                    result = [i.strip() for i in r if i]
                                    if result:
                                        if not r[0].startswith('---'):
                                            discipline = ''.join(result[0].split('\n')).upper().strip()
                                            if not discipline.startswith('КЛАССНЫЙ') and not discipline.startswith('КЛ.') and not discipline.startswith('КЛАС.'):
                                                if group.lower() == search.lower() and discipline.strip():
                                                    # header = result[-1]
                                                    # months[int(dd)] = months.get(int(dd), {})
                                                    # months[int(dd)][discipline + '_' + header] = months[int(dd)].get(
                                                    #     discipline + '_' + header, 0) + 1
                                                    months[int(dd)] = months.get(int(dd), 0) + 2




            except BaseException:
                print('load_files schedule')

        progress_callback.emit(100)

        return months

    def update_table(self, data=None):
        if hasattr(self, 'table'):
            self.table.deleteLater()

        self.table = QtWidgets.QTableWidget()
        self.table.horizontalHeader().setVisible(False)
        self.table.verticalHeader().setVisible(False)
        self.table.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.ResizeMode.Stretch)

        if data:
            data = data.copy()
            n = len(data)

            n_top = int(len(data) ** (1 / 2) + 1)
            n_bottom = int(len(data) ** (1 / 2))

            self.table.setColumnCount(n_top)
            self.table.setRowCount(n_bottom * 2)

            for i in range(n_top):
                for j in range(0, n_bottom * 2, 2):

                    if data:
                        k, v = data.popitem()
                        title = QTableWidgetItem(str(k))
                        value = QTableWidgetItem(str(v))


                        self.table.setItem(j, i, title)
                        self.table.setItem(j+1, i, value)


                        # self.table.setItem(j, i, QTableWidgetItem())

        self.table_layout.addWidget(self.table)

    def click_load_data_button(self):
        self.is_load = True
        self.close()

    def closeEvent(self, arg__1):
        self.is_stop = True
        super().closeEvent(arg__1)




class BaseTable:
    SHOW_COUNT_ROWS = 15

    def save_table(self):
        try:
            MANAGER_STUDENTS.save_students()
            CM.IS_CHANGE = False
        except BaseException:
            self.parent.get_down_message(self.tr('НЕ удалось сохранить файл'))
        else:
            self.parent.get_down_message(self.tr('Успешное сохранение'))


class AbsenceTab(QtWidgets.QWidget, BaseTable):
    def __init__(self, parent):
        super(AbsenceTab, self).__init__()
        self.setStyleSheet('''QRadioButton::indicator {
        width: 15px;
        height: 15px;
        }
        QRadioButton::indicator::unchecked {
            background-color: gray;
        }
        #is_sick_rb::indicator:unchecked:hover {
            background-color: rgb(51, 153, 51);
        }
        #is_absen_rb::indicator:unchecked:hover {
            background-color: rgb(255, 165, 0);
        }
        QRadioButton::indicator:unchecked:pressed {
            background-color: gray;
        }
        #is_sick_rb::indicator::checked {
            background-color: rgb(51, 153, 51);
            border: 2px solid rgb(51, 153, 51);
        }
        #is_absen_rb::indicator::checked {
            background-color: rgb(255, 165, 0);
            border: 2px solid rgb(255, 165, 0);
        }
        
        ''')

        self.parent = parent
        self.setObjectName("F6")
        self.verticalLayout = QtWidgets.QVBoxLayout(self)
        self.verticalLayout.setObjectName("verticalLayout")

        self.verticalLayout_3 = QtWidgets.QVBoxLayout()
        self.verticalLayout_3.setObjectName("verticalLayout_3")
        self.frame = QtWidgets.QFrame(self)
        self.frame.setMinimumSize(QtCore.QSize(0, 80))
        self.frame.setFrameShape(QtWidgets.QFrame.Shape.StyledPanel)
        self.frame.setFrameShadow(QtWidgets.QFrame.Shadow.Raised)
        self.frame.setObjectName("frame")
        self.horizontalLayout_3 = QtWidgets.QHBoxLayout(self.frame)
        self.horizontalLayout_3.setObjectName("horizontalLayout_3")
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setSizeConstraint(QtWidgets.QLayout.SizeConstraint.SetMinimumSize)
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.label_2 = QtWidgets.QLabel(self.frame)
        self.label_2.setMinimumSize(QtCore.QSize(250, 0))
        self.label_2.setObjectName("label_2")
        self.label_2.setFont(QtGui.QFont('Calibri', 15))
        self.horizontalLayout.addWidget(self.label_2)
        self.horizontalLayout_3.addLayout(self.horizontalLayout)
        self.wLayout = QtWidgets.QVBoxLayout()
        self.statustic1 = QtWidgets.QLabel(self.frame)
        self.statustic2 = QtWidgets.QLabel(self.frame)

        self.wLayout.addWidget(self.statustic1)
        self.wLayout.addWidget(self.statustic2)
        self.horizontalLayout_3.addLayout(self.wLayout)

        spacerItem = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Policy.Expanding,
                                           QtWidgets.QSizePolicy.Policy.Minimum)
        self.horizontalLayout_3.addItem(spacerItem)
        self.horizontalLayout_2 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_2.setContentsMargins(7, -1, -1, -1)
        self.horizontalLayout_2.setObjectName("horizontalLayout_2")
        self.horizontalLayout_4 = QtWidgets.QVBoxLayout()
        self.horizontalLayout_4.setContentsMargins(7, -1, -1, -1)
        self.horizontalLayout_4.setObjectName("horizontalLayout_4")
        self.is_sick_rb = QtWidgets.QRadioButton(self.frame)

        self.is_sick_rb.setObjectName("is_sick_rb")
        self.horizontalLayout_4.addWidget(self.is_sick_rb)
        self.radioButton_2 = QtWidgets.QRadioButton(self.frame)
        self.radioButton_2.setObjectName("is_absen_rb")
        self.radioButton_2.setChecked(True)
        self.horizontalLayout_4.addWidget(self.radioButton_2)

        self.double_mod = QtWidgets.QCheckBox()
        self.horizontalLayout_2.addWidget(self.double_mod)
        self.double_mod.hide()

        # ________________________________________
        self.load_schedule_push = Push(self.frame, 40, 40, 5,
                                       icon_path=os.path.join(CONTENT_PATH, 'media', 'buttons', 'load_schedule.svg'))
        self.load_schedule_push.setObjectName("load_schedule_push")
        self.horizontalLayout_2.addWidget(self.load_schedule_push)

        self.set_size_posetiv_font_push = Push(self.frame, 40, 40, 5,
                                               icon_path=os.path.join(CONTENT_PATH, 'media', 'posetive.svg'))

        self.horizontalLayout_2.addWidget(self.set_size_posetiv_font_push)

        self.set_size_negativ_font_push = Push(self.frame, 40, 40, 5,
                                               icon_path=os.path.join(CONTENT_PATH, 'media', 'negative.svg'))
        self.horizontalLayout_2.addWidget(self.set_size_negativ_font_push)

        # _________________________________________

        self.horizontalLayout_2.addLayout(self.horizontalLayout_4)

        self.save_table_push = Push(self.frame, 40, 40, 5,
                                    icon_path=os.path.join(CONTENT_PATH, 'media', 'buttons', 'save.svg'))
        self.save_table_push.setObjectName("save_table_push")
        self.horizontalLayout_2.addWidget(self.save_table_push)
        self.game_over_push = Push(self.frame, 40, 40, 5,
                                   icon_path=os.path.join(CONTENT_PATH, 'media', 'game_over.svg'))
        self.game_over_push.setObjectName("game_over_push")
        self.horizontalLayout_2.addWidget(self.game_over_push)
        self.save_to_exel_push = Push(self.frame, 55, 55, 5,
                                      icon_path=os.path.join(CONTENT_PATH, 'media', 'save_exel.svg'))
        self.save_to_exel_push.setObjectName("save_to_exel_push")
        self.horizontalLayout_2.addWidget(self.save_to_exel_push)
        self.horizontalLayout_3.addLayout(self.horizontalLayout_2)

        # _______________Кнопки после таблицы___________________________________

        self.properties_table_view_horizontalLayout = QtWidgets.QHBoxLayout()

        self.pages_message_1 = QtWidgets.QLabel("Страница ")
        self.properties_table_view_horizontalLayout.addWidget(self.pages_message_1)
        self.current_page = QtWidgets.QLabel("0")
        self.properties_table_view_horizontalLayout.addWidget(self.current_page)
        self.pages_message_2 = QtWidgets.QLabel("из")
        self.properties_table_view_horizontalLayout.addWidget(self.pages_message_2)
        self.total_pages = QtWidgets.QLabel("1")
        self.properties_table_view_horizontalLayout.addWidget(self.total_pages)
        spacerItem = QtWidgets.QSpacerItem(100, 1, QtWidgets.QSizePolicy.Policy.Expanding,
                                           QtWidgets.QSizePolicy.Policy.Minimum)
        self.properties_table_view_horizontalLayout.addItem(spacerItem)

        self.section_first = QtWidgets.QHBoxLayout()
        self.labe1_show_items_layout = QtWidgets.QLabel("Выводить ")
        self.labe1_show_items_layout.setAlignment(QtCore.Qt.AlignmentFlag.AlignRight)
        self.section_first.addWidget(self.labe1_show_items_layout)

        self.count_rows_edit = QtWidgets.QLineEdit()
        self.count_rows_edit.setText(str(self.SHOW_COUNT_ROWS))
        self.count_rows_edit.setValidator(QtGui.QIntValidator(bottom=0))
        self.count_rows_edit.setMaximumWidth(100)
        self.section_first.addWidget(self.count_rows_edit)

        self.section_second = QtWidgets.QHBoxLayout()
        self.all_rows_checkbox = QtWidgets.QCheckBox()
        self.section_second.addWidget(self.all_rows_checkbox)
        self.labe2_show_items_layout = QtWidgets.QLabel("Вывести всё ")
        self.section_second.addWidget(self.labe2_show_items_layout)

        self.properties_table_view_horizontalLayout.addLayout(self.section_first)
        self.properties_table_view_horizontalLayout.addLayout(self.section_second)

        self.arrow_narrow_top_push = Push(self.properties_table_view_horizontalLayout, 25, 25, 5,
                                          icon_path=os.path.join(CONTENT_PATH, 'media', 'arrow_narrow_top.svg'))
        self.arrow_narrow_top_push.setObjectName("arrow_narrow_top_push")
        self.properties_table_view_horizontalLayout.addWidget(self.arrow_narrow_top_push)

        self.arrow_narrow_down_push = Push(self.properties_table_view_horizontalLayout, 25, 25, 5,
                                           icon_path=os.path.join(CONTENT_PATH, 'media', 'arrow_narrow_down.svg'))
        self.arrow_narrow_down_push.setObjectName("arrow_narrow_down_push")
        self.properties_table_view_horizontalLayout.addWidget(self.arrow_narrow_down_push)

        self.update_table_push = Push(self.properties_table_view_horizontalLayout, 25, 25, 5,
                                      icon_path=os.path.join(CONTENT_PATH, 'media', 'update.svg'))
        self.update_table_push.setObjectName("update_table_push")
        self.properties_table_view_horizontalLayout.addWidget(self.update_table_push)

        self.verticalLayout_3.addWidget(self.frame)

        self.verticalLayout.addLayout(self.verticalLayout_3)
        self.table_layout = QtWidgets.QVBoxLayout(self)
        self.verticalLayout.addLayout(self.table_layout)
        self.verticalLayout.addLayout(self.properties_table_view_horizontalLayout)

        self.retranslateUi()
        self.add_function()

        self.is_click_end_period = 0
        self.number_page = 0

    def add_function(self):
        self.save_to_exel_push.clicked.connect(self.save_to_exel)
        self.save_table_push.clicked.connect(self.save_table)
        self.game_over_push.clicked.connect(self.click_end_period)
        self.all_rows_checkbox.stateChanged.connect(self.click_all_rows_checkbox)
        self.count_rows_edit.textChanged.connect(self.change_count_rows_edit)
        self.arrow_narrow_top_push.clicked.connect(self.click_arrow_narrow_top_push)
        self.arrow_narrow_down_push.clicked.connect(self.click_arrow_narrow_down_push)
        self.update_table_push.clicked.connect(self.click_update_table_push)
        self.load_schedule_push.clicked.connect(self.click_load_shedule_push)

    def index(self, selected, deselected):
        for ix in selected.indexes():
            if self.tableWidget.item(ix.row(), ix.column()):
                print(ix.row(), ix.column())

    # def keyPressEvent(self, e):
    #     pass

    def retranslateUi(self):

        self.is_sick_rb.setText(self.tr("ПОУВ"))
        self.radioButton_2.setText(self.tr("НЕУВ"))

        if USER_MANAGER.user:
            self.label_2.setText(self.tr('Добро пожаловать, ') + str(USER_MANAGER.user.username) + '!')
        else:
            self.label_2.setText(self.tr('Добро пожаловать, Noname!'))

        self.set_size_posetiv_font_push.setToolTip(self.tr('Увеличить размер текста'))
        self.set_size_negativ_font_push.setToolTip(self.tr('Уменьшить размер текста'))
        self.save_table_push.setToolTip(self.tr('Сохранить изменения'))
        self.game_over_push.setToolTip(self.tr('Завершить текущий месяц'))
        self.save_to_exel_push.setToolTip(self.tr('Сохранить в EXEL'))
        self.double_mod.setText('2X')

        self.is_sick_rb.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.radioButton_2.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.label_2.setFont(QtGui.QFont(CM.NAME_FONT, 17 + CM.ADD_FONT_SIZE))

        self.statustic1.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))
        self.double_mod.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))
        self.statustic2.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))

        if hasattr(self, 'tableWidget'):
            self.tableWidget.retranslateUi()
            self.update_statistics()

    def click_load_shedule_push(self):
        b = ScheduleLoadWindow(self)
        b.exec()
        data = b.result
        b.deleteLater()

        if b.is_load:
            for k, v in data.items():
                MANAGER_STUDENTS.add_hours_by_day(k, v)
            self.tableWidget.update_hours_by_days()
        del b



    def update_statistics(self):
        statistics = MANAGER_STUDENTS.get_statistics()
        if statistics.get('man_hours') > 0:
            self.statustic1.setText(
                self.tr('Чел.час = ') + f'{str(statistics.get("man_hours"))} \t' +
                self.tr(' Посещ.Кач. = ') + f'{str(round(statistics.get("quality_attendance") * 100, 2))}%')
            self.statustic2.setText(
                self.tr('Посещ.Общ. = ') + f'{str(round(statistics.get("total_attendance") * 100, 2))}% \t' +
                self.tr('Прогул 1 студ. = ') + f'{str(round(statistics.get("absences_by_student", 1)))}')
        else:
            self.statustic1.clear()
            self.statustic2.clear()

    def save_to_exel(self):
        # options = QtWidgets.QFileDialog.Options()
        # options |= QtWidgets.QFileDialog.Option.DontUseNativeDialog

        file_name, _ = QtWidgets.QFileDialog.getSaveFileName(self, self.tr("Выберите путь и имя файла для сохранения."),
                                                             f'{"_".join(["П", str(MANAGER_STUDENTS.MONTHS[MANAGER_STUDENTS.period[0] - 1]), str(MANAGER_STUDENTS.period[1]), str(USER_MANAGER.user.username)])}.xlsx',
                                                             self.tr("Книга Excel (*.xlsx)"))

        if file_name:
            try:
                path = MANAGER_STUDENTS.save_f6(file_name)
            except BaseException as f:
                QtWidgets.QMessageBox.critical(self, self.tr('Ошибка сохранения'),
                                               self.tr('Файл не был сохранен. Повторите попытку'))
            else:
                self.parent.get_down_message(self.tr('Файл успешно сохранен в директории: ') + path, time=3000)

    def set_size_font(self, is_posetiv=True):
        # print(id(self.tableWidget))
        if is_posetiv and self.tableWidget.mod_size < 25:
            self.tableWidget.update_table_students(size=self.tableWidget.mod_size + 2)
        elif is_posetiv and self.tableWidget.mod_size > 25:
            self.parent.profile.cod = [1, 4]
        elif self.tableWidget.mod_size < -7:
            self.parent.profile.cod = [1, 3]
        elif self.tableWidget.mod_size > -8:
            self.tableWidget.update_table_students(size=self.tableWidget.mod_size - 2)

    def clicked_table(self, tablewidget):
        CM.IS_CHANGE = True
        if self.parent.group.indexOf(self.parent.students) == -1 and len(MANAGER_STUDENTS.students) != 0:
            self.parent.group.insertTab(2, self.parent.students, self.tr('Cтуденты'))

        item = tablewidget.currentItem()

        if not item is None:
            hours = item.text()
            try:

                if item.row() == 1 and 2 <= item.column() <= 32:
                    if item.text() == '' or item.text() == ' ':
                        if MANAGER_STUDENTS.days.get(item.column() - 1):
                            MANAGER_STUDENTS.days[item.column() - 1] = 0
                        item.setText('')
                    elif item.text().isnumeric() and 0 <= int(item.text()) <= 10:
                        MANAGER_STUDENTS.add_hours_by_day(item.column() - 1, int(item.text()))
                    else:
                        item.setText(str(MANAGER_STUDENTS.days.get(item.column() - 1, '')))
                    item.setFont(QtGui.QFont('Calibri', 14 + tablewidget.mod_size))
                    tablewidget.update_hours_day()

                elif 32 >= item.column() >= 2 and len(MANAGER_STUDENTS.students) + 2 >= item.row() >= 2:
                    if item.text() == '' or item.text() == ' ':
                        if item.column() - 1 in MANAGER_STUDENTS.students[item.row() - 3].absence_days:
                            del MANAGER_STUDENTS.students[item.row() - 3].absence_days[item.column() - 1]
                        if item.column() - 1 in MANAGER_STUDENTS.students[item.row() - 3].sick_days:
                            del MANAGER_STUDENTS.students[item.row() - 3].sick_days[item.column() - 1]
                        item.setBackground(QtGui.QColor(255, 255, 255))

                    elif hours.isnumeric() and 1 <= (
                            int(hours) * 2 if self.double_mod.isChecked() else int(hours)) <= 10:
                        hours = int(hours)
                        if self.is_sick_rb.isChecked():
                            tablewidget.add_hours_in_table(item.row(), item.column(), hours, type_day='s')
                        else:
                            tablewidget.add_hours_in_table(item.row(), item.column(), hours, type_day='a')


                    else:
                        if MANAGER_STUDENTS.students[item.row() - 3].sick_days.get(item.column() - 1, ''):
                            item.setText(str(
                                MANAGER_STUDENTS.students[item.row() - 3].sick_days.get(item.column() - 1, '')))
                        elif MANAGER_STUDENTS.students[item.row() - 3].absence_days.get(item.column() - 1, ''):
                            item.setText(str(
                                MANAGER_STUDENTS.students[item.row() - 3].absence_days.get(item.column() - 1, '')))
                        else:
                            item.setText('')
                    tablewidget.update_statistics_student(item.row())

                elif item.column() == 1:
                    if item.row() == len(MANAGER_STUDENTS.students) + 3:
                        try:
                            MANAGER_STUDENTS.CLASS_STUDENT.check_fio(item.text())
                        except BaseException:
                            item.setBackground(QtGui.QColor(255, 0, 0))
                        else:
                            self.parent.profile.cod = [0, 1]
                            if item.text().lower() == 'чайковский пётр ильич' or item.text().lower() == 'чайковский петр ильич':
                                self.parent.profile.cod = [2, 2]
                            if len(MANAGER_STUDENTS.students) < 30:
                                MANAGER_STUDENTS.add_student(MANAGER_STUDENTS.CLASS_STUDENT(item.text()))
                                for i in range(2, 32 + 1):
                                    if not i - 1 in MANAGER_STUDENTS.days:
                                        tablewidget.setItem(tablewidget.rowCount() - 2, i,
                                                            QTableWidgetItem(""))
                                        tablewidget.item(tablewidget.rowCount() - 2, i).setBackground(
                                            QtGui.QColor(220, 220, 220))
                                        tablewidget.item(tablewidget.rowCount() - 2, i).setFlags(
                                            QtCore.Qt.ItemFlag.ItemIsEnabled)
                            tablewidget.update_table_students()
                            if hasattr(self.parent, 'marks'):
                                self.parent.marks.tableWidget_3.update_table_students()
                            if hasattr(self.parent, 'students'):
                                self.parent.students.update_list_students()

            except BaseException as f:
                print(f)
            self.update_statistics()

    def click_end_period(self):
        self.is_click_end_period = 1
        self.parent.profile.cod = [0, 2]
        i = 1
        message = QtWidgets.QMessageBox.question(self, self.tr('Завершение текущего месяца'),
                                                 self.tr(
                                                     'После завершения рабочего месяца, таблицы будут сохранены в архиве, но изменение данных в них будет уже не доступно. Продолжить?'),
                                                 QtWidgets.QMessageBox.StandardButton.No | QtWidgets.QMessageBox.StandardButton.Yes)
        if message.Yes == message:
            self.save_table()
            MANAGER_STUDENTS.save_students()

            try:
                MANAGER_STUDENTS.push_archive()  #
            except FileExistsError:

                massage = QtWidgets.QMessageBox.question(self, self.tr('Ошибка'),
                                                         self.tr(
                                                             'Файл уже существует. Заменить его?\nДа - заменит старый файл на новый\nНет - файл будет помещен рядом'),
                                                         QtWidgets.QMessageBox.StandardButton.No | QtWidgets.QMessageBox.StandardButton.Yes)

                if massage == massage.Yes:
                    MANAGER_STUDENTS.del_file_archive(MANAGER_STUDENTS.create_archive_file_name())
                    MANAGER_STUDENTS.push_archive()
                else:
                    MANAGER_STUDENTS.push_archive(auto=True)

            except FileNotFoundError:
                print('Отстутствует активный файл')

            MANAGER_STUDENTS.create_new_table()
            # self.update_statistics()
            # self.parent.marks.update_statistics_2()
            # self.tableWidget.update_table_students() было
            self.init_table_absence()
            self.parent.marks.tableWidget_3.update_table_students()
            self.parent.archive.init_archive()
            self.parent.profile.cod = [3, 2]

        if self.parent.group.indexOf(self.parent.archive) == -1 and len(self.parent.archive.files_archive) != 0:
            self.parent.group.insertTab(3, self.parent.archive, self.tr('Архив'))

    def change_count_rows_edit(self):
        try:
            count = int(self.count_rows_edit.text())
        except ValueError:
            count = self.SHOW_COUNT_ROWS

        self.number_page = 0
        self.show_select_row(end=count)

        self.total_pages.setText(
            str(math.ceil((self.tableWidget.rowCount() - 3) / (int(self.count_rows_edit.text())))))
        self.current_page.setText(str(1))

    def click_arrow_narrow_top_push(self):
        if self.number_page > 0:
            self.number_page -= 1
        start = self.number_page * (int(self.count_rows_edit.text()))
        end = (self.number_page + 1) * (int(self.count_rows_edit.text()))
        self.show_select_row(start=start, end=end)
        self.current_page.setText(str(self.number_page + 1))
        self.total_pages.setText(
            str(math.ceil((self.tableWidget.rowCount() - 3) / (int(self.count_rows_edit.text())))))

    def click_arrow_narrow_down_push(self):
        if self.number_page + 1 < math.ceil((self.tableWidget.rowCount() - 3) / (int(self.count_rows_edit.text()))):
            self.number_page += 1
        start = self.number_page * (int(self.count_rows_edit.text()))
        end = (self.number_page + 1) * (int(self.count_rows_edit.text()))
        self.show_select_row(start=start, end=end)
        self.current_page.setText(str(self.number_page + 1))
        self.total_pages.setText(
            str(math.ceil((self.tableWidget.rowCount() - 3) / (int(self.count_rows_edit.text())))))

    def click_all_rows_checkbox(self):
        if self.all_rows_checkbox.isChecked():
            self.show_select_row()
            self.count_rows_edit.setReadOnly(True)
            self.pages_message_1.hide()
            self.pages_message_2.hide()
            self.total_pages.hide()
            self.current_page.hide()

        else:
            self.show_select_row(end=int(self.count_rows_edit.text()))
            self.count_rows_edit.setReadOnly(False)

            self.pages_message_1.show()
            self.pages_message_2.show()
            self.total_pages.show()
            self.current_page.show()

            self.total_pages.setText(
                str(math.ceil((self.tableWidget.rowCount() - 3) / (int(self.count_rows_edit.text())))))
            self.current_page.setText(str(1))

    def click_update_table_push(self):
        self.init_table_absence()

    def init_table_absence(self, only_show=False):
        if hasattr(self, 'tableWidget'):
            # self.tableWidget.hide()
            self.table_layout.removeWidget(self.tableWidget)
            self.tableWidget.deleteLater()
            self.tableWidget = None

        self.tableWidget = TableAbsence()

        self.tableWidget.setStyleSheet("""QTableWidget {border: none;}""")
        self.tableWidget.setObjectName("tableView")

        self.table_layout.addWidget(self.tableWidget)
        # self.verticalLayout.addLayout(self.table_layout)
        self.verticalLayout.addItem(
            QtWidgets.QSpacerItem(0, 0, QtWidgets.QSizePolicy.Policy.Expanding, QtWidgets.QSizePolicy.Policy.Minimum))

        self.set_size_negativ_font_push.clicked.connect(lambda: self.set_size_font(False))
        self.set_size_posetiv_font_push.clicked.connect(lambda: self.set_size_font())
        self.tableWidget.update_table_students()

        self.update_statistics()

        self.click_all_rows_checkbox()

        if not only_show:
            self.tableWidget.cellChanged.connect(lambda: self.clicked_table(self.tableWidget))
            self.tableWidget.currentCellChanged.connect(self.cellPressed)
            self.update_table_push.hide()
            # self.tableWidget.selectionModel().selectionChanged.connect(self.index)
        else:
            self.update_table_push.hide()

    def show_select_row(self, start=0, end=0, count_head_rows=3):
        if not end:
            end = self.tableWidget.rowCount()
        for row in range(0, self.tableWidget.rowCount()):
            if start + count_head_rows <= row < end + count_head_rows or row < count_head_rows:
                self.tableWidget.setRowHidden(row, False)
            else:
                self.tableWidget.setRowHidden(row, True)

    def cellPressed(self, row, column):
        if self.tableWidget.hasFocus():
            if 32 >= column >= 2 and len(
                    MANAGER_STUDENTS.students) + 2 >= row >= 3 and column - 1 in MANAGER_STUDENTS.days:
                try:
                    if hasattr(self.tableWidget, 'old_item1'):
                        self.tableWidget.item(self.tableWidget.old_item1[0], 0).setBackground(
                            QtGui.QColor(255, 255, 255))
                        self.tableWidget.item(2, self.tableWidget.old_item1[1]).setBackground(
                            QtGui.QColor(255, 255, 255))
                        self.tableWidget.old_item1 = (row, column)
                        self.tableWidget.item(self.tableWidget.old_item1[0], 0)
                    else:
                        self.tableWidget.old_item1 = (row, column)
                    self.tableWidget.item(row, 0).setBackground(QtGui.QColor(102, 102, 102))
                    self.tableWidget.item(2, column).setBackground(QtGui.QColor(102, 102, 102))

                except BaseException as f:
                    print(repr(f))


class MarksTab(QtWidgets.QWidget, BaseTable):
    def __init__(self, parent):
        super(MarksTab, self).__init__()
        self.parent = parent
        self.setObjectName("marks")
        self.verticalLayout_25 = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_25.setObjectName("verticalLayout_25")
        self.verticalLayout_24 = QtWidgets.QVBoxLayout()
        self.verticalLayout_24.setObjectName("verticalLayout_24")
        self.frame_3 = QtWidgets.QFrame(self)
        self.frame_3.setMinimumSize(QtCore.QSize(0, 80))
        self.frame_3.setFrameShape(QtWidgets.QFrame.Shape.StyledPanel)
        self.frame_3.setFrameShadow(QtWidgets.QFrame.Shadow.Raised)
        self.frame_3.setObjectName("frame_3")
        self.horizontalLayout_17 = QtWidgets.QHBoxLayout(self.frame_3)
        self.horizontalLayout_17.setObjectName("horizontalLayout_17")

        self.wLayout2 = QtWidgets.QVBoxLayout()
        self.statustic3 = QtWidgets.QLabel(self.frame_3)
        self.statustic4 = QtWidgets.QLabel(self.frame_3)
        self.wLayout2.addWidget(self.statustic3)
        self.wLayout2.addWidget(self.statustic4)

        self.horizontalLayout_17.addLayout(self.wLayout2)

        spacerItem1 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        self.horizontalLayout_17.addItem(spacerItem1)
        self.horizontalLayout_19 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_19.setContentsMargins(7, -1, -1, -1)
        self.horizontalLayout_19.setObjectName("horizontalLayout_19")
        self.pushButton_9 = Push(self.frame_3, 40, 40, 5,
                                 icon_path=os.path.join(CONTENT_PATH, 'media', 'buttons', 'save.svg'))
        self.pushButton_9.setObjectName("pushButton_9")

        # ___________________________________________________________________
        self.set_size_posetiv_font_push = Push(self.frame_3, 40, 40, 5,
                                               icon_path=os.path.join(CONTENT_PATH, 'media', 'posetive.svg'))
        self.horizontalLayout_19.addWidget(self.set_size_posetiv_font_push)

        self.set_size_negativ_font_push = Push(self.frame_3, 40, 40, 5,
                                               icon_path=os.path.join(CONTENT_PATH, 'media', 'negative.svg'))
        self.horizontalLayout_19.addWidget(self.set_size_negativ_font_push)

        # ___________________________________________________________________
        self.horizontalLayout_19.addWidget(self.pushButton_9)
        self.save_to_exel_marks_push = Push(self.frame_3, 55, 55, 5,
                                            icon_path=os.path.join(CONTENT_PATH, 'media', 'save_exel.svg'))
        self.save_to_exel_marks_push.setObjectName("pushButton_10")
        self.horizontalLayout_19.addWidget(self.save_to_exel_marks_push)
        self.horizontalLayout_17.addLayout(self.horizontalLayout_19)

        # ____________________После таблицы_____________________________
        self.properties_table_view_horizontalLayout = QtWidgets.QHBoxLayout()

        self.pages_message_1 = QtWidgets.QLabel("Страница ")
        self.properties_table_view_horizontalLayout.addWidget(self.pages_message_1)
        self.current_page = QtWidgets.QLabel("0")
        self.properties_table_view_horizontalLayout.addWidget(self.current_page)
        self.pages_message_2 = QtWidgets.QLabel("из")
        self.properties_table_view_horizontalLayout.addWidget(self.pages_message_2)
        self.total_pages = QtWidgets.QLabel("1")
        self.properties_table_view_horizontalLayout.addWidget(self.total_pages)
        spacerItem = QtWidgets.QSpacerItem(100, 1, QtWidgets.QSizePolicy.Policy.Expanding,
                                           QtWidgets.QSizePolicy.Policy.Minimum)
        self.properties_table_view_horizontalLayout.addItem(spacerItem)

        self.section_first = QtWidgets.QHBoxLayout()
        self.labe1_show_items_layout = QtWidgets.QLabel("Выводить ")
        self.labe1_show_items_layout.setAlignment(QtCore.Qt.AlignmentFlag.AlignRight)
        self.section_first.addWidget(self.labe1_show_items_layout)

        self.count_rows_edit = QtWidgets.QLineEdit()
        self.count_rows_edit.setText(str(self.SHOW_COUNT_ROWS))
        self.count_rows_edit.setValidator(QtGui.QIntValidator(bottom=0))
        self.count_rows_edit.setMaximumWidth(100)
        self.section_first.addWidget(self.count_rows_edit)

        self.section_second = QtWidgets.QHBoxLayout()
        self.all_rows_checkbox = QtWidgets.QCheckBox()
        self.section_second.addWidget(self.all_rows_checkbox)
        self.labe2_show_items_layout = QtWidgets.QLabel("Вывести всё ")
        self.section_second.addWidget(self.labe2_show_items_layout)

        self.properties_table_view_horizontalLayout.addLayout(self.section_first)
        self.properties_table_view_horizontalLayout.addLayout(self.section_second)

        self.arrow_narrow_top_push = Push(self.properties_table_view_horizontalLayout, 25, 25, 5,
                                          icon_path=os.path.join(CONTENT_PATH, 'media', 'arrow_narrow_top.svg'))
        self.arrow_narrow_top_push.setObjectName("arrow_narrow_top_push")
        self.properties_table_view_horizontalLayout.addWidget(self.arrow_narrow_top_push)

        self.arrow_narrow_down_push = Push(self.properties_table_view_horizontalLayout, 25, 25, 5,
                                           icon_path=os.path.join(CONTENT_PATH, 'media', 'arrow_narrow_down.svg'))
        self.arrow_narrow_down_push.setObjectName("arrow_narrow_down_push")
        self.properties_table_view_horizontalLayout.addWidget(self.arrow_narrow_down_push)

        # self.update_table_push = Push(self.properties_table_view_horizontalLayout, 25, 25, 5,
        #                               icon_path=os.path.join('media', 'update.svg'))
        # self.update_table_push.setObjectName("update_table_push")
        # self.properties_table_view_horizontalLayout.addWidget(self.update_table_push)

        self.verticalLayout_24.addWidget(self.frame_3)
        self.verticalLayout_25.addLayout(self.verticalLayout_24)
        self.table_layout = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_25.addLayout(self.table_layout)
        self.verticalLayout_25.addLayout(self.properties_table_view_horizontalLayout)

        self.retranslateUi()
        self.add_function()

        self.number_page = 0

    def retranslateUi(self):
        self.set_size_posetiv_font_push.setToolTip(self.tr('Увеличить размер текста'))
        self.set_size_negativ_font_push.setToolTip(self.tr('Уменьшить размер текста'))
        self.pushButton_9.setToolTip(self.tr('Сохранить изменения'))
        self.save_to_exel_marks_push.setToolTip(self.tr('Сохранить в EXEL'))

        if hasattr(self, 'tableWidget_3'):
            self.tableWidget_3.retranslateUi()
            self.update_statistics_2()

        self.statustic3.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))
        self.statustic4.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))

    def add_function(self):
        self.pushButton_9.clicked.connect(self.save_table)
        self.save_to_exel_marks_push.clicked.connect(self.save_to_exel_marks)

        self.all_rows_checkbox.stateChanged.connect(self.click_all_rows_checkbox)
        self.count_rows_edit.textChanged.connect(self.change_count_rows_edit)
        self.arrow_narrow_top_push.clicked.connect(self.click_arrow_narrow_top_push)
        self.arrow_narrow_down_push.clicked.connect(self.click_arrow_narrow_down_push)

    def update_statistics_2(self):
        statistics = MANAGER_STUDENTS.get_statistics_marks()
        if statistics.get('is_ready') == 1:
            self.statustic3.setText(
                self.tr('Кол-во студ. им-х 2 = ') + str(statistics.get("heaving_2")) + ' | ' + self.tr(
                    'Кол-во студ. им-х одну 3 = ') + str(statistics.get("heaving_one_3")) + ' \t| ' + self.tr(
                    'Успеваемость общая = ') + str(
                    round(statistics.get("total_academic_performance", 0) * 100, 2)) + '%')
            self.statustic4.setText(
                self.tr('Кол-во студ. им-х 5 = ') + str(statistics.get("heaving_5")) + ' | ' + self.tr(
                    'Кол-во студ. им-х 4 и 5 = ') + str(statistics.get("heaving_4_and_5")) + ' \t| ' + self.tr(
                    'Успеваемость качественная = ') + str(
                    round(statistics.get("quality_academic_performance", 0) * 100, 2)) + '%')
        else:
            self.statustic3.clear()
            self.statustic4.clear()

    def change_count_rows_edit(self):
        try:
            count = int(self.count_rows_edit.text())
        except ValueError:
            count = self.SHOW_COUNT_ROWS

        self.number_page = 0
        self.show_select_row(end=count)

        self.total_pages.setText(
            str(math.ceil((self.tableWidget_3.rowCount() - 3) / (int(self.count_rows_edit.text())))))
        self.current_page.setText(str(1))

    def show_select_row(self, start=0, end=0, count_head_rows=3):
        if not end:
            end = self.tableWidget_3.rowCount()
        for row in range(0, self.tableWidget_3.rowCount()):
            if start + count_head_rows <= row < end + count_head_rows or row < count_head_rows:
                self.tableWidget_3.setRowHidden(row, False)
            else:
                self.tableWidget_3.setRowHidden(row, True)

    def click_arrow_narrow_top_push(self):
        if self.number_page > 0:
            self.number_page -= 1
        start = self.number_page * (int(self.count_rows_edit.text()))
        end = (self.number_page + 1) * (int(self.count_rows_edit.text()))
        self.show_select_row(start=start, end=end)
        self.current_page.setText(str(self.number_page + 1))
        self.total_pages.setText(
            str(math.ceil((self.tableWidget_3.rowCount() - 3) / (int(self.count_rows_edit.text())))))

    def click_arrow_narrow_down_push(self):
        if self.number_page + 1 < math.ceil((self.tableWidget_3.rowCount() - 3) / (int(self.count_rows_edit.text()))):
            self.number_page += 1
        start = self.number_page * (int(self.count_rows_edit.text()))
        end = (self.number_page + 1) * (int(self.count_rows_edit.text()))
        self.show_select_row(start=start, end=end)
        self.current_page.setText(str(self.number_page + 1))
        self.total_pages.setText(
            str(math.ceil((self.tableWidget_3.rowCount() - 3) / (int(self.count_rows_edit.text())))))

    def click_all_rows_checkbox(self):
        if self.all_rows_checkbox.isChecked():
            self.show_select_row()
            self.count_rows_edit.setReadOnly(True)
            self.pages_message_1.hide()
            self.pages_message_2.hide()
            self.total_pages.hide()
            self.current_page.hide()

        else:
            self.show_select_row(end=int(self.count_rows_edit.text()))
            self.count_rows_edit.setReadOnly(False)

            self.pages_message_1.show()
            self.pages_message_2.show()
            self.total_pages.show()
            self.current_page.show()

            self.total_pages.setText(
                str(math.ceil((self.tableWidget_3.rowCount() - 3) / (int(self.count_rows_edit.text())))))
            self.current_page.setText(str(1))

    def init_table_marks(self, only_show=False):
        if hasattr(self, 'tableWidget_3'):
            # self.tableWidget.hide()
            self.table_layout.removeWidget(self.tableWidget_3)
            self.tableWidget_3.deleteLater()
            self.tableWidget_3 = None

        self.tableWidget_3 = TableMarks(only_show=only_show)
        self.tableWidget_3.setObjectName("tableWidget_3")
        self.table_layout.addWidget(self.tableWidget_3)
        self.verticalLayout_25.addItem(
            QtWidgets.QSpacerItem(0, 0, QtWidgets.QSizePolicy.Policy.Expanding, QtWidgets.QSizePolicy.Policy.Minimum))

        self.set_size_negativ_font_push.clicked.connect(lambda: self.set_size_font(False))
        self.set_size_posetiv_font_push.clicked.connect(lambda: self.set_size_font())
        self.tableWidget_3.update_table_students()

        self.click_all_rows_checkbox()

        if not only_show:
            self.tableWidget_3.currentCellChanged.connect(self.cellPressed)
            self.tableWidget_3.cellChanged.connect(self.clicked_table_marks)

    def clicked_table_marks(self):
        CM.IS_CHANGE = True
        item = self.tableWidget_3.currentItem()
        if not item is None:
            if item.column() == 1:
                if item.row() == len(MANAGER_STUDENTS.students) + 3:
                    try:
                        MANAGER_STUDENTS.CLASS_STUDENT.check_fio(item.text())
                    except BaseException:
                        item.setBackground(QtGui.QColor(255, 0, 0))
                    else:
                        try:
                            MANAGER_STUDENTS.add_student(MANAGER_STUDENTS.CLASS_STUDENT(item.text()))
                        except BaseException as f:
                            print(f)
                        else:
                            self.tableWidget_3.update_table_students()
                            if hasattr(self.parent, 'F6'):
                                self.parent.F6.tableWidget.update_table_students()
                            if hasattr(self.parent, 'students'):
                                self.parent.students.update_list_students()
            elif 32 >= item.column() >= 2 and len(MANAGER_STUDENTS.students) + 2 >= item.row() >= 3:
                if item.text().isnumeric() and 5 >= int(item.text()) >= 2:
                    self.tableWidget_3.add_mark_table(item.row(), item.column(), item.text())
                else:
                    item.setText('')
                    self.tableWidget_3.del_mark_table(item.row(), item.column())

            elif item.row() == 1 and item.column() >= 2:
                self.tableWidget_3.add_couples(item.column() - 1, couple=item.text())
            elif item.row() == 2 and item.column() >= 2:
                if MANAGER_STUDENTS.CLASS_STUDENT.is_valud_fio(item.text()):
                    self.tableWidget_3.add_couples(item.column() - 1, fio=item.text().title())
                    item.setText(item.text().title())
                    item.setBackground(QtGui.QColor(255, 255, 255))
                else:
                    item.setBackground(QtGui.QColor(255, 0, 0))
        self.update_statistics_2()

    def set_size_font(self, is_posetiv=True):
        if is_posetiv and self.tableWidget_3.mod_size < 25:
            self.tableWidget_3.update_table_students(size=self.tableWidget_3.mod_size + 2)
        elif self.tableWidget_3.mod_size > -8:
            self.tableWidget_3.update_table_students(size=self.tableWidget_3.mod_size - 2)

    def save_to_exel_marks(self):

        file_name, _ = QtWidgets.QFileDialog.getSaveFileName(self, self.tr("Выберите путь и имя файла для сохранения."),
                                                             f'{"_".join(["О", str(MANAGER_STUDENTS.MONTHS[MANAGER_STUDENTS.period[0] - 1]), str(MANAGER_STUDENTS.period[1]), str(USER_MANAGER.user.username)])}.xlsx',
                                                             self.tr("Книга Excel (*.xlsx)"))

        if file_name:
            try:
                path = MANAGER_STUDENTS.save_f6_marks(file_name)
            except BaseException as f:
                QtWidgets.QMessageBox.critical(self,
                                               self.tr('Ошибка сохранения', 'Файл не был сохранен. Повторите попытку'))
            else:
                self.parent.get_down_message(self.tr('Файл успешно сохранен в директории: ') + path, time=3000)

    def cellPressed(self, row, column):
        if self.tableWidget_3.hasFocus():
            if 32 >= column >= 2 and row >= 3 and len(
                    MANAGER_STUDENTS.students) + 2 >= row >= 3:
                try:
                    if hasattr(self.tableWidget_3, 'old_item1'):
                        self.tableWidget_3.item(self.tableWidget_3.old_item1[0], 0).setBackground(
                            QtGui.QColor(255, 255, 255))
                        self.tableWidget_3.item(1, self.tableWidget_3.old_item1[1]).setBackground(
                            QtGui.QColor(255, 255, 255))
                        self.tableWidget_3.old_item1 = (row, column)
                        self.tableWidget_3.item(self.tableWidget_3.old_item1[0], 0)
                    else:
                        self.tableWidget_3.old_item1 = (row, column)
                    self.tableWidget_3.item(row, 0).setBackground(QtGui.QColor(102, 102, 102))
                    self.tableWidget_3.item(1, column).setBackground(QtGui.QColor(102, 102, 102))

                except BaseException as f:
                    print(repr(f))


class StudentsTab(QtWidgets.QWidget):
    def __init__(self, parent):
        super(StudentsTab, self).__init__(parent=parent)
        self.parent = parent
        self.setStyleSheet("""
                        #del_push {
                           font-size: 16px;
                           color: black;
                           background-color: red;
                           border: none;
                           padding: 5px;
                           color: white;
                           }
                        #del_push:hover {
                           font: 18px;
                           border: 2px solid #990000;
                           }
                        #save_push {
                           font-size: 16px;
                           color: black;
                           background-color: green;
                           border: none;
                           padding: 5px;
                           color: white;
                           }
                    
                        #save_push:hover {
                           font: 18px;
                           border: 2px solid #009900;
                           }

                        QLineEdit {
                           border: none;
                           color:white;
                           padding-left: 15px;
                           padding-right: 15px;
                           font: 16px;
                           background: black;
                           }
                           """)
        self.setObjectName("students")
        self.horizontalLayout_6 = QtWidgets.QHBoxLayout(self)
        self.horizontalLayout_6.setObjectName("horizontalLayout_6")
        self.verticalLayout_6 = QtWidgets.QVBoxLayout()
        self.verticalLayout_6.setContentsMargins(-1, -1, 0, -1)
        self.verticalLayout_6.setObjectName("verticalLayout_6")
        self.label = QtWidgets.QLabel(self)

        self.label.setObjectName("label")
        self.verticalLayout_6.addWidget(self.label)
        self.listWidget = QtWidgets.QListWidget(self)
        self.listWidget.setObjectName("listWidget")
        self.listWidget.setStyleSheet('''#listWidget {margin: auto; border: none;} ''')
        self.verticalLayout_6.addWidget(self.listWidget)
        self.horizontalLayout_6.addLayout(self.verticalLayout_6)
        self.verticalLayout_4 = QtWidgets.QVBoxLayout()
        self.verticalLayout_4.setObjectName("verticalLayout_4")
        self.verticalLayout_5 = QtWidgets.QVBoxLayout()
        self.verticalLayout_5.setObjectName("verticalLayout_5")
        self.fio_label = QtWidgets.QLabel(self)

        self.fio_label.setObjectName("fio_label")
        self.verticalLayout_5.addWidget(self.fio_label)
        self.fio_edit = QtWidgets.QLineEdit(self)
        self.fio_edit.setMinimumSize(QtCore.QSize(0, 30))
        self.fio_edit.setMaximumSize(QtCore.QSize(468, 16777215))
        self.fio_edit.setObjectName("fio_edit")
        self.fio_edit.setEnabled(0)
        self.verticalLayout_5.addWidget(self.fio_edit)
        self.message_students = QtWidgets.QTextBrowser(self)
        self.message_students.setFixedSize(400, 400)
        self.message_students.setStyleSheet(
            'border: none; color: red; font: 14px; background-color: rgba(249, 248, 244, 0);')
        self.verticalLayout_5.addWidget(self.message_students)

        spacerItem1 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.verticalLayout_5.addItem(spacerItem1)
        self.horizontalLayout_5 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_5.setObjectName("horizontalLayout_5")
        self.del_push = QtWidgets.QPushButton(self)
        self.del_push.setObjectName("del_push")
        self.horizontalLayout_5.addWidget(self.del_push)
        self.save_push = QtWidgets.QPushButton(self)
        self.save_push.setObjectName("save_push")
        self.horizontalLayout_5.addWidget(self.save_push)
        self.verticalLayout_5.addLayout(self.horizontalLayout_5)
        self.verticalLayout_4.addLayout(self.verticalLayout_5)
        self.horizontalLayout_6.addLayout(self.verticalLayout_4)

        self.add_function()
        self.retranslateUi()

        self.del_push.setEnabled(0)
        self.save_push.setEnabled(0)

    def retranslateUi(self):
        self.fio_label.setText(self.tr("ФИО"))
        self.del_push.setText(self.tr("Удалить"))
        self.save_push.setText(self.tr("Сохранить"))
        self.label.setText(self.tr("Список студентов"))

        self.fio_label.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))
        self.del_push.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))
        self.save_push.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))
        self.label.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 16))
        self.listWidget.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))

    def add_function(self):
        self.save_push.clicked.connect(self.save_student)
        self.del_push.clicked.connect(self.del_student)

        self.listWidget.itemClicked.connect(lambda: self.click_list(self.listWidget))

    def update_list_students(self):
        self.listWidget.clear()
        for indx, student in enumerate(MANAGER_STUDENTS.students):
            self.listWidget.addItem(QtWidgets.QListWidgetItem(f'{str(indx + 1)}. {student.fio}'))

    def save_student(self):
        self.message_students.setText('')
        if self.listWidget.currentItem():
            if MANAGER_STUDENTS.CLASS_STUDENT.is_valud_fio(self.fio_edit.text()):
                MANAGER_STUDENTS.students[
                    self.listWidget.row(self.listWidget.currentItem())].fio = self.fio_edit.text()
                self.update_list_students()
                if hasattr(self.parent, 'marks'):
                    self.parent.marks.tableWidget_3.update_table_students()
                if hasattr(self.parent, 'F6'):
                    self.parent.F6.tableWidget.update_table_students()
                MANAGER_STUDENTS.save_students()
                self.parent.get_down_message(self.tr('Изменения сохранены'))
                self.del_push.setEnabled(0)
                self.save_push.setEnabled(0)
            else:
                self.message_students.setText(self.tr(
                    'ФИО должно состоять только из букв и быть из 3 частей, каждая из которых не менее 2 символов'))

    def del_student(self):
        if self.listWidget.currentItem():
            message_main = QtWidgets.QMessageBox().question(self, self.tr("Удаление студента"),
                                                            self.tr(
                                                                'Вы точно хотите навсегда удалить студента ') + " ".join(
                                                                self.listWidget.currentItem().text().split()[1:]) + '?',
                                                            QtWidgets.QMessageBox.StandardButton.Yes | QtWidgets.QMessageBox.StandardButton.No)
            if message_main == message_main.Yes:
                MANAGER_STUDENTS.remove_student_id(self.listWidget.row(self.listWidget.currentItem()))
                if hasattr(self.parent, 'marks'):
                    self.parent.marks.tableWidget_3.update_table_students()
                if hasattr(self.parent, 'F6'):
                    self.parent.F6.tableWidget.update_table_students()
                self.listWidget.removeItemWidget(self.listWidget.currentItem())
                self.update_list_students()
                USER_MANAGER.save_users()
                self.parent.get_down_message(self.tr('Студент отчислен'))
                self.del_push.setEnabled(0)
                self.save_push.setEnabled(0)

        if len(MANAGER_STUDENTS.students) == 0:
            self.parent.group.removeTab(self.parent.group.indexOf(self))
            self.update_list_students()

    def click_list(self, listwidget):
        self.del_push.setEnabled(1)
        self.save_push.setEnabled(1)
        self.fio_edit.setText(' '.join(listwidget.currentItem().text().split()[1:]))
        self.fio_edit.setEnabled(1)


class ArchiveTab(QtWidgets.QWidget):
    def __init__(self, parent):
        super(ArchiveTab, self).__init__(parent=parent)
        self.parent = parent
        self.setObjectName("archive")
        self.setStyleSheet("""
                                #del_file_push {
                                   font-size: 16px;
                                   color: black;
                                   background-color: red;
                                   border: none;
                                   padding: 5px;
                                   color: white;
                                   }
                                #del_file_push:hover {
                                   font: 18px;
                                   border: 2px solid #990000;
                                   }
                                #load_file_push {
                                   font-size: 16px;
                                   color: black;
                                   background-color: green;
                                   border: none;
                                   padding: 5px;
                                   color: white;
                                   }

                                #load_file_push:hover {
                                   font: 18px;
                                   border: 2px solid #009900;
                                   }
                                   """)
        self.verticalLayout_11 = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_11.setObjectName("verticalLayout_11")
        self.label_4 = QtWidgets.QLabel(self)

        self.label_4.setObjectName("label_4")
        self.verticalLayout_11.addWidget(self.label_4)
        self.list_archive = QtWidgets.QListWidget(self)
        self.list_archive.setObjectName("list_archive")

        self.verticalLayout_11.addWidget(self.list_archive)

        self.horizontalLayout_9 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_9.setContentsMargins(-1, 10, 0, -1)
        self.horizontalLayout_9.setObjectName("horizontalLayout_9")
        # spacerItem3 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Policy.Expanding,
        #                                     QtWidgets.QSizePolicy.Policy.Minimum)
        # self.horizontalLayout_9.addItem(spacerItem3)
        self.del_file_push = QtWidgets.QPushButton(self)
        self.del_file_push.setObjectName("del_file_push")
        self.horizontalLayout_9.addWidget(self.del_file_push)
        self.load_file_push = QtWidgets.QPushButton(self)
        self.load_file_push.setObjectName("load_file_push")
        self.horizontalLayout_9.addWidget(self.load_file_push)
        self.verticalLayout_11.addLayout(self.horizontalLayout_9)
        self.logout_archive_push = Push(self.parent.F6.frame, 40, 40, 5, tool_tip=self.tr('Выйти из архива'),
                                        icon_path=os.path.join(CONTENT_PATH, 'media', 'undo.svg'))
        self.logout_archive_push.hide()
        self.logout_archive_push.setObjectName("logout_archive_push")

        self.parent.F6.horizontalLayout_2.addWidget(self.logout_archive_push)
        self.add_function()
        self.retranslateUi()

        self.is_comebake = 0
        self.is_delite_file = 0

    def retranslateUi(self):
        self.label_4.setText(self.tr("Доступные файлы"))
        self.del_file_push.setText(self.tr("Удалить"))
        self.load_file_push.setText(self.tr("Активировать"))

        self.list_archive.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))
        self.label_4.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 16))
        self.del_file_push.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))
        self.load_file_push.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))

    def add_function(self):
        self.del_file_push.clicked.connect(self.clicked_del_file_push)
        self.load_file_push.clicked.connect(self.clicked_load_file_push)

        self.logout_archive_push.clicked.connect(self.logout_archive)

    def init_archive(self):
        path, files = MANAGER_STUDENTS.init_archive()
        self.list_archive.clear()
        for i in files:
            self.list_archive.addItem(QtWidgets.QListWidgetItem(i))
        self.path_archive = path
        self.files_archive = files

    def clicked_load_file_push(self):
        item = self.list_archive.currentItem()
        if item:
            try:
                self.parent.init_students_manager(path=os.path.join('archive', item.text()), only_show=True)
            except BaseException as f:
                QtWidgets.QMessageBox.critical(self, self.tr('Просмотр архивного файла'),
                                               self.tr(
                                                   f'При попытки  загрузки файла произошла ошибка. Попробуйте загрузить его заново. Если проблема не исчезнет, то скорее всего файл поврежден или удален.'),
                                               QtWidgets.QMessageBox.StandardButton.Ok)
            else:
                self.logout_archive_push.show()

                self.parent.F6.save_table_push.hide()
                self.parent.F6.game_over_push.hide()
                self.parent.F6.radioButton_2.hide()
                self.parent.marks.pushButton_9.hide()
                self.parent.F6.is_sick_rb.hide()

                self.parent.group.removeTab(self.parent.group.indexOf(self.parent.students))
                self.parent.group.removeTab(self.parent.group.indexOf(self.parent.settings))

                QtWidgets.QMessageBox.information(self, self.tr('Просмотр архивного файла'),
                                                  self.tr('Файла успешно загружен. Вы можете перейти к просмотру.'),
                                                  QtWidgets.QMessageBox.StandardButton.Ok)

    def clicked_del_file_push(self):
        item = self.list_archive.currentItem()
        if item:
            self.del_file_archive(item.text())
            self.is_delite_file = 1
            self.parent.profile.cod = [2, 4]

    def del_file_archive(self, file_name):
        if file_name in self.files_archive:
            message = QtWidgets.QMessageBox.question(self.parent, self.tr('Удаление архивного файла'),
                                                     self.tr('Вы точно хотите удалить файл: ') +
                                                     file_name +
                                                     self.tr('?\n После этого он будет удален навсегда'),
                                                     QtWidgets.QMessageBox.StandardButton.Yes | QtWidgets.QMessageBox.StandardButton.No)
            if message == message.Yes:
                try:
                    os.remove(os.path.join(self.path_archive, file_name))
                except BaseException:
                    pass
                finally:
                    self.init_archive()

            if len(self.parent.archive.files_archive) == 0 and self.parent.group.indexOf(self.parent.archive) != -1:
                self.parent.group.removeTab(self.parent.group.indexOf(self.parent.archive))

    def logout_archive(self):
        self.is_comebake = 1
        self.parent.profile.cod = [3, 1]
        self.parent.init_students_manager(path='students.json')
        self.logout_archive_push.hide()

        self.parent.F6.save_table_push.show()
        self.parent.F6.game_over_push.show()
        self.parent.marks.pushButton_9.show()
        self.parent.F6.radioButton_2.show()
        self.parent.F6.is_sick_rb.show()

        self.parent.group.insertTab(2, self.parent.students, self.tr('Cтуденты'))
        self.parent.group.insertTab(4, self.parent.settings, self.tr('Настройки'))


class SettingsTab(QtWidgets.QWidget):
    def __init__(self, parent):
        super(SettingsTab, self).__init__(parent=parent)
        self.parent = parent
        QtWidgets.QWidget()
        self.setObjectName("settings")
        self.setStyleSheet(
            '''#scrollAreaWidgetContents {background-color:white}'''
        )
        self.verticalLayout_10 = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_10.setObjectName("verticalLayout_10")
        self.scrollArea = QtWidgets.QScrollArea(self)
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        # self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 950, 712))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")

        self.main_layout = QtWidgets.QHBoxLayout(self.scrollAreaWidgetContents)
        self.left_layout = QtWidgets.QVBoxLayout()
        self.left_layout.setObjectName("left_layout")
        self.setings1_label = QtWidgets.QLabel()
        self.setings1_label.setObjectName("setings1_label")
        self.left_layout.addWidget(self.setings1_label)
        self.show_current_month_link_button = QtWidgets.QCommandLinkButton()
        self.show_current_month_link_button.setTabletTracking(True)
        self.show_current_month_link_button.setObjectName("show_current_month_link_button")
        self.left_layout.addWidget(self.show_current_month_link_button)
        self.del_work_day_link_button = QtWidgets.QCommandLinkButton()
        self.del_work_day_link_button.setCheckable(False)
        self.del_work_day_link_button.setObjectName("del_work_day_link_button")
        self.del_work_day_link_button.hide()
        self.left_layout.addWidget(self.del_work_day_link_button)
        self.restart_weekend_link_button = QtWidgets.QCommandLinkButton()
        self.restart_weekend_link_button.setTabletTracking(True)
        self.restart_weekend_link_button.setObjectName("restart_weekend_link_button")
        self.left_layout.addWidget(self.restart_weekend_link_button)
        self.show_weekend_link_button = QtWidgets.QCommandLinkButton()
        self.show_weekend_link_button.setTabletTracking(True)
        self.show_weekend_link_button.setObjectName("show_weekend_link_button")
        self.left_layout.addWidget(self.show_weekend_link_button)
        line = QtWidgets.QFrame()
        line.setFrameShape(QtWidgets.QFrame.Shape.HLine)
        line.setFrameShadow(QtWidgets.QFrame.Shadow.Sunken)
        line.setObjectName("line")
        self.left_layout.addWidget(line)
        self.on_off_table_marks_label = QtWidgets.QLabel()
        self.on_off_table_marks_label.setObjectName("on_off_table_marks_label")
        self.left_layout.addWidget(self.on_off_table_marks_label)
        self.on_off_table_marks_link_button = QtWidgets.QCommandLinkButton()
        self.on_off_table_marks_link_button.setTabletTracking(True)
        self.on_off_table_marks_link_button.setObjectName("on_off_table_marks_link_button")
        self.left_layout.addWidget(self.on_off_table_marks_link_button)
        self.on_off_statistics_link_button = QtWidgets.QCommandLinkButton()
        self.on_off_statistics_link_button.setTabletTracking(True)
        self.on_off_statistics_link_button.setObjectName("on_off_table_marks_link_button")
        self.left_layout.addWidget(self.on_off_statistics_link_button)
        line = QtWidgets.QFrame()
        line.setFrameShape(QtWidgets.QFrame.Shape.HLine)
        line.setFrameShadow(QtWidgets.QFrame.Shadow.Sunken)
        line.setObjectName("line")
        self.left_layout.addWidget(line)
        self.clear_table_label = QtWidgets.QLabel()
        self.clear_table_label.setObjectName("clear_table_label")
        self.left_layout.addWidget(self.clear_table_label)
        self.clear_table_abcense_link_button = QtWidgets.QCommandLinkButton()
        self.clear_table_abcense_link_button.setCheckable(False)
        self.clear_table_abcense_link_button.setObjectName("clear_table_abcense_link_button")
        self.left_layout.addWidget(self.clear_table_abcense_link_button)
        self.clear_table_marks_link_button = QtWidgets.QCommandLinkButton()
        self.clear_table_marks_link_button.setTabletTracking(True)
        self.clear_table_marks_link_button.setObjectName("clear_table_marks_link_button")
        self.left_layout.addWidget(self.clear_table_marks_link_button)
        line = QtWidgets.QFrame(self.scrollAreaWidgetContents)
        line.setFrameShape(QtWidgets.QFrame.Shape.HLine)
        line.setFrameShadow(QtWidgets.QFrame.Shadow.Sunken)
        line.setObjectName("line")
        self.left_layout.addWidget(line)
        self.set_data_table_label = QtWidgets.QLabel()
        self.set_data_table_label.setObjectName("set_data_table_label")
        self.left_layout.addWidget(self.set_data_table_label)
        self.set_data_table_link_button = QtWidgets.QCommandLinkButton()
        self.set_data_table_link_button.setCheckable(False)
        self.set_data_table_link_button.setObjectName("set_data_table_link_button")
        self.left_layout.addWidget(self.set_data_table_link_button)
        line = QtWidgets.QFrame()
        line.setFrameShape(QtWidgets.QFrame.Shape.HLine)
        line.setFrameShadow(QtWidgets.QFrame.Shadow.Sunken)
        line.setObjectName("line")
        self.left_layout.addWidget(line)
        self.language_label = QtWidgets.QLabel()
        self.language_label.setObjectName("language_label")
        self.left_layout.addWidget(self.language_label)
        self.set_language_link_button = QtWidgets.QCommandLinkButton()
        self.set_language_link_button.setCheckable(False)
        self.set_language_link_button.setObjectName("set_language_link_button")
        self.left_layout.addWidget(self.set_language_link_button)

        spacerItem2 = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                            QtWidgets.QSizePolicy.Policy.Expanding)
        self.left_layout.addItem(spacerItem2)

        self.right_layout = QtWidgets.QVBoxLayout()
        self.right_layout.setObjectName("right_layout")
        self.setings1_label_test = QtWidgets.QLabel()
        self.setings1_label_test.setObjectName("setings1_label")

        self.right_layout.addWidget(self.setings1_label_test)
        self.show_current_month_link_button_test = QtWidgets.QCommandLinkButton()
        self.show_current_month_link_button_test.setTabletTracking(True)
        self.show_current_month_link_button_test.setObjectName("show_current_month_link_button_test")
        self.right_layout.addWidget(self.show_current_month_link_button_test)
        self.restart_weekend_link_button_test = QtWidgets.QCommandLinkButton()
        self.restart_weekend_link_button_test.setTabletTracking(True)
        self.restart_weekend_link_button_test.setObjectName("restart_weekend_link_button_test")
        self.right_layout.addWidget(self.restart_weekend_link_button_test)
        self.show_weekend_link_button_test = QtWidgets.QCommandLinkButton()
        self.show_weekend_link_button_test.setTabletTracking(True)
        self.show_weekend_link_button_test.setObjectName("show_weekend_link_button_test")
        self.right_layout.addWidget(self.show_weekend_link_button_test)

        # self.show_current_month_link_button_test.hide()
        self.restart_weekend_link_button_test.hide()
        self.show_weekend_link_button_test.hide()


        line = QtWidgets.QFrame()
        line.setFrameShape(QtWidgets.QFrame.Shape.HLine)
        line.setFrameShadow(QtWidgets.QFrame.Shadow.Sunken)
        line.setObjectName("line")
        self.right_layout.addWidget(line)
        self.right_layout.addItem(QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                                        QtWidgets.QSizePolicy.Policy.Expanding))

        self.main_layout.addLayout(self.left_layout)
        self.main_layout.addLayout(self.right_layout)

        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.verticalLayout_10.addWidget(self.scrollArea)

        self.add_function()
        self.retranslateUi()

        self.is_add_work_day = 0
        self.is_del_work_day = 0

    def retranslateUi(self):

        self.setings1_label.setText(self.tr("Рабочие/нерабочие дни"))
        self.show_current_month_link_button.setText(self.tr("Текущий месяц"))
        self.set_data_table_link_button.setText(self.tr("Изменить месяц и год"))
        self.del_work_day_link_button.setText(self.tr("Удалить рабочий день"))
        self.restart_weekend_link_button.setText(self.tr('Сброс праздников'))
        self.show_weekend_link_button.setText(self.tr('Праздничные дни'))

        self.set_data_table_label.setText(self.tr('Период'))
        self.language_label.setText(self.tr('Язык'))
        self.set_language_link_button.setText(self.tr('Язык'))
        self.clear_table_label.setText(self.tr('Очистка таблицы'))
        self.clear_table_marks_link_button.setText(self.tr('Очистить таблицу оценок'))
        self.clear_table_abcense_link_button.setText(self.tr('Очистить таблицу прогулов'))
        self.on_off_table_marks_label.setText(self.tr('Отображение вкладок'))
        self.on_off_table_marks_link_button.setText(self.tr('Включить/выключить таблицу оценок'))
        self.on_off_statistics_link_button.setText(self.tr('Включить/выключить статистику'))

        self.setings1_label.setFont(QtGui.QFont(CM.NAME_FONT, 16 + CM.ADD_FONT_SIZE))
        self.show_current_month_link_button.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.restart_weekend_link_button.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.show_weekend_link_button.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.set_data_table_link_button.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.del_work_day_link_button.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.set_data_table_label.setFont(QtGui.QFont(CM.NAME_FONT, 16 + CM.ADD_FONT_SIZE))
        self.set_data_table_label.setFont(QtGui.QFont(CM.NAME_FONT, 16 + CM.ADD_FONT_SIZE))
        self.set_language_link_button.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.clear_table_label.setFont(QtGui.QFont(CM.NAME_FONT, 16 + CM.ADD_FONT_SIZE))
        self.clear_table_marks_link_button.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.language_label.setFont(QtGui.QFont(CM.NAME_FONT, 16 + CM.ADD_FONT_SIZE))
        self.clear_table_abcense_link_button.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.on_off_table_marks_label.setFont(QtGui.QFont(CM.NAME_FONT, 16 + CM.ADD_FONT_SIZE))
        self.on_off_table_marks_link_button.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.on_off_statistics_link_button.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))

    def add_function(self):
        self.show_current_month_link_button.clicked.connect(self.show_current_month)
        self.del_work_day_link_button.clicked.connect(self.show_academic_year)
        self.set_data_table_link_button.clicked.connect(self.set_data_table)
        self.clear_table_marks_link_button.clicked.connect(self.clear_table_marks)
        self.clear_table_abcense_link_button.clicked.connect(self.clear_table_abcense)
        self.on_off_table_marks_link_button.clicked.connect(self.on_off_table)
        self.on_off_statistics_link_button.clicked.connect(self.on_off_statistics)
        self.set_language_link_button.clicked.connect(self.set_language)
        self.restart_weekend_link_button.clicked.connect(self.restart_weekend)
        self.show_weekend_link_button.clicked.connect(self.show_year)
        self.show_current_month_link_button_test.clicked.connect(self.show_weekend)

    def set_language(self):
        global set_language
        set_language()

    def restart_weekend(self):
        message = QtWidgets.QMessageBox.question(self, self.tr('Сброс'),
                                                 self.tr('Вы точно хотите "сбросить" все праздники?'),
                                                 QtWidgets.QMessageBox.StandardButton.Yes,
                                                 QtWidgets.QMessageBox.StandardButton.No)
        if message == QtWidgets.QMessageBox.StandardButton.Yes:
            USER_MANAGER.user.restart_happy_day()
            MANAGER_STUDENTS.days = MANAGER_STUDENTS.generate_work_days(month=MANAGER_STUDENTS.period[0],
                                                                        year=MANAGER_STUDENTS.period[1],
                                                                        happy_days=USER_MANAGER.user.happy_days)
            self.parent.F6.tableWidget.update_table_students()
            self.parent.F6.update_statistics()

    def show_weekend(self):
        self.list_weekend = QtWidgets.QTextBrowser()
        self.list_weekend.setWindowIcon(QtGui.QIcon(os.path.join(CONTENT_PATH, 'media', 'logo.svg')))
        self.list_weekend.setWindowTitle(self.tr('Список выходных дней'))
        self.list_weekend.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))

        resalt = '<table>'
        for k, v in USER_MANAGER.user.happy_days.items():
            resalt += f'<tr><td>{MANAGER_STUDENTS.MONTHS[int(k) - 1]}</td><td>{str(v)[1:-1]}</td></tr>'
        resalt += '</table>'
        self.list_weekend.setText(resalt)

        self.list_weekend.show()

    def show_current_month(self):
        month = MonthCalendarView(MANAGER_STUDENTS, parent=self.parent)
        month.show()
        month.exec()
        try:
            MANAGER_STUDENTS.user.save_happy_days()
        except BaseException as f:
            self.parent.get_down_message(self.tr('НЕ удалось сохранить файл'))
        else:
            try:
                MANAGER_STUDENTS.save_students()
                CM.IS_CHANGE = False
            except BaseException:
                self.parent.get_down_message(self.tr('НЕ удалось сохранить файл'))
            else:
                self.parent.get_down_message(self.tr('Успешное сохранение'))

        month.deleteLater()

    def show_year(self):
        window = QtWidgets.QDialog()#parent=self.parent)
        window.setWindowTitle(self.tr('Праздничные дни'))
        window.setWindowIcon(QtGui.QIcon(os.path.join(CONTENT_PATH, 'media', 'logo.svg')))
        window.setModal(True)
        centerLayout = QtWidgets.QVBoxLayout(window)
        scrollArea = QtWidgets.QScrollArea(window)
        scrollArea.setWidgetResizable(True)
        scrollAreaWidgetContents = QtWidgets.QWidget()
        layout = QtWidgets.QVBoxLayout(scrollAreaWidgetContents)
        month = MANAGER_STUDENTS.period[0] - 1
        for i in range(4):
            layout_row = QtWidgets.QHBoxLayout()
            for j in range(3):
                layout_row.addWidget(YearCalendarView(MANAGER_STUDENTS, period=[month % 12 + 1, MANAGER_STUDENTS.period[1] + month // 12], parent=self.parent, is_only_happy_days=True))
                month += 1
            layout.addLayout(layout_row)
        scrollArea.setWidget(scrollAreaWidgetContents)
        centerLayout.addWidget(scrollArea)
        scrollAreaWidgetContents.show()
        window.show()
        window.exec()
        try:
            MANAGER_STUDENTS.user.save_happy_days()
        except BaseException as f:
            self.parent.get_down_message(self.tr('НЕ удалось сохранить файл'))
        else:
            try:
                MANAGER_STUDENTS.save_students()
                CM.IS_CHANGE = False
            except BaseException:
                self.parent.get_down_message(self.tr('НЕ удалось сохранить файл'))
            else:
                self.parent.get_down_message(self.tr('Успешное сохранение'))

        window.deleteLater()

    def show_academic_year(self):
        pass

    def add_work_day(self):
        deal = SettingsWindows(self)
        deal.show()
        deal.exec()
        if not (deal.result is None):
            MANAGER_STUDENTS.on_day(deal.result)
            self.parent.F6.tableWidget.update_table_students()
            self.parent.F6.update_statistics()
            self.is_add_work_day = 1
        self.parent.profile.cod = [1, 0]

    def del_work_day(self):
        deal = SettingsWindows(self)
        deal.show()
        deal.exec()
        if not (deal.result is None):
            MANAGER_STUDENTS.off_day(deal.result)

            self.parent.F6.tableWidget.update_table_students()
            self.parent.F6.update_statistics()
            self.is_del_work_day = 1
        self.parent.profile.cod = [1, 1]

    def set_data_table(self):
        try:
            s = SettingsData(self)
            s.show()
            s.exec()
        except BaseException as f:
            print(repr(f))
        else:
            if s.status == 'finished':
                self.parent.init_students_manager(period=tuple(map(lambda x: int(x), s.dateEdit.text().split('.'))))

    def del_table_marks(self):
        MANAGER_STUDENTS.clear_marks()

    def clear_table_marks(self):
        message = QtWidgets.QMessageBox.question(self, self.tr('Очистка таблицы оценок'),
                                                 self.tr('Вы точно хотите очистить таблицу оценок?'),
                                                 QtWidgets.QMessageBox.StandardButton.Yes,
                                                 QtWidgets.QMessageBox.StandardButton.No)
        if message == QtWidgets.QMessageBox.StandardButton.Yes:
            MANAGER_STUDENTS.clear_marks()
            MANAGER_STUDENTS.save_students()
            if hasattr(self.parent, 'marks'):
                self.parent.marks.tableWidget_3.update_table_students()

    def clear_table_abcense(self):
        message = QtWidgets.QMessageBox.question(self, self.tr('Очистка таблицы прогулов'),
                                                 self.tr('Вы точно хотите очистить таблицу прогулов?'),
                                                 QtWidgets.QMessageBox.StandardButton.Yes,
                                                 QtWidgets.QMessageBox.StandardButton.No)
        if message == QtWidgets.QMessageBox.StandardButton.Yes:
            MANAGER_STUDENTS.clear_absences()
            MANAGER_STUDENTS.save_students()
            if hasattr(self.parent, 'F6'):
                self.parent.F6.tableWidget.update_table_students()

    def on_off_table(self):
        if self.parent.group.indexOf(self.parent.marks) != -1:
            self.parent.group.removeTab(self.parent.group.indexOf(self.parent.marks))
            USER_MANAGER.user.parametrs['table_marks'] = False
        else:
            self.parent.group.insertTab(1, self.parent.marks, self.tr('Оценки'))
            USER_MANAGER.user.parametrs['table_marks'] = True
        USER_MANAGER.user.save_user()

    def on_off_statistics(self):
        if self.parent.group.indexOf(self.parent.statistics) != -1:
            self.parent.group.removeTab(self.parent.group.indexOf(self.parent.statistics))
            USER_MANAGER.user.parametrs['statistic'] = False
        else:
            self.parent.group.insertTab(5, self.parent.statistics, self.tr('Статистика'))
            USER_MANAGER.user.parametrs['statistic'] = True
        USER_MANAGER.user.save_user()


class ProfileTab(QtWidgets.QWidget):
    def __init__(self, parent):
        super(ProfileTab, self).__init__(parent=parent)
        self.setStyleSheet("""
                                QLineEdit, #username_edit{
                                   border: none;
                                   color:white;
                                   font: 16px;
                                   padding-left: 15px;
                                   background: black;
                                  
                                   }
                                #scrollAreaWidgetContents {background-color:white}
                                   """)
        self.parent = parent
        self.setObjectName("profile")
        self.base_layout = QtWidgets.QVBoxLayout(self)
        self.scrollArea = QtWidgets.QScrollArea(self)
        self.scrollArea.setWidgetResizable(True)

        self.scrollArea.setObjectName("scrollArea")

        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        # self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 25, 712))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")

        self.center_layout = QtWidgets.QHBoxLayout(self.scrollAreaWidgetContents)
        self.center_layout.setObjectName("center_layout")
        self.center_layout.addItem(QtWidgets.QSpacerItem(5000, 0, QtWidgets.QSizePolicy.Policy.Maximum,
                                                         QtWidgets.QSizePolicy.Policy.Expanding))
        self.left_layout = QtWidgets.QVBoxLayout()
        self.left_layout.setSizeConstraint(QtWidgets.QLayout.SizeConstraint.SetMinimumSize)
        self.left_layout.setObjectName("left_layout")

        self.information_user_layout = QtWidgets.QVBoxLayout()

        self.accaunt_label = QtWidgets.QLabel()
        self.accaunt_label.setObjectName("accaunt_label")
        self.information_user_layout.addWidget(self.accaunt_label)

        self.username_label = QtWidgets.QLabel()
        self.username_label.setObjectName("username_label")
        self.information_user_layout.addWidget(self.username_label)

        self.username_edit = QtWidgets.QLabel()
        self.username_edit.setFixedSize(QtCore.QSize(400, 30))
        self.username_edit.setObjectName("username_edit")
        self.information_user_layout.addWidget(self.username_edit)

        self.fio_user_label = QtWidgets.QLabel()
        self.fio_user_label.setObjectName("fio_user_label")
        self.information_user_layout.addWidget(self.fio_user_label)

        self.fio_user_edit = QtWidgets.QLineEdit()
        self.fio_user_edit.setFixedSize(QtCore.QSize(400, 30))
        self.fio_user_edit.setObjectName("fio_user_edit")
        self.information_user_layout.addWidget(self.fio_user_edit)

        self.teamleader_label = QtWidgets.QLabel()
        self.teamleader_label.setObjectName("teamleader_label")
        self.information_user_layout.addWidget(self.teamleader_label)

        self.fio_teamleader_edit = QtWidgets.QLineEdit()
        self.fio_teamleader_edit.setFixedSize(QtCore.QSize(400, 30))
        self.fio_teamleader_edit.setObjectName("fio_teamleader_edit")
        self.information_user_layout.addWidget(self.fio_teamleader_edit)

        self.group_label = QtWidgets.QLabel()
        self.group_label.setObjectName("group_label")
        self.information_user_layout.addWidget(self.group_label)

        self.group_edit = QtWidgets.QLineEdit()
        self.group_edit.setFixedSize(QtCore.QSize(400, 30))
        self.group_edit.setObjectName("group_edit")
        self.information_user_layout.addWidget(self.group_edit)

        self.specialization_label = QtWidgets.QLabel()
        self.specialization_label.setObjectName("specialization_label")
        self.information_user_layout.addWidget(self.specialization_label)

        self.specialization_edit = QtWidgets.QLineEdit()
        self.specialization_edit.setFixedSize(QtCore.QSize(400, 30))
        self.specialization_edit.setObjectName("specialization_edit")
        self.information_user_layout.addWidget(self.specialization_edit)

        self.message_profile = QtWidgets.QTextBrowser()
        self.message_profile.setFixedWidth(400)
        self.message_profile.setStyleSheet(
            'border: none; color: red; font: 14px; background-color: rgba(249, 248, 244, 0);')
        self.information_user_layout.addWidget(self.message_profile)

        self.left_layout.addLayout(self.information_user_layout)
        # self.left_layout.addItem(QtWidgets.QSpacerItem(20, 2000, QtWidgets.QSizePolicy.Policy.Maximum,
        #                                                QtWidgets.QSizePolicy.Policy.Expanding))

        self.security_level_layout = QtWidgets.QVBoxLayout()
        # self.security_level_layout.addItem(QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Maximum,
        #                                                      QtWidgets.QSizePolicy.Policy.Expanding))
        self.security_leve_label = QtWidgets.QLabel()
        self.security_leve_label.setObjectName("security_leve_label")
        self.security_level_layout.addWidget(self.security_leve_label)

        self.security_level_photo_layout = QtWidgets.QVBoxLayout()
        self.update_photo_security_level()
        self.security_level_layout.addLayout(self.security_level_photo_layout)
        # self.security_level_layout.addItem(QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Maximum,
        #                                                          QtWidgets.QSizePolicy.Policy.Expanding))
        self.security_level_slider = QtWidgets.QSlider(QtCore.Qt.Orientation.Horizontal)
        self.security_level_slider.setTickPosition(QtWidgets.QSlider.TickPosition.TicksBelow)
        self.security_level_slider.setTickInterval(50)
        self.security_level_slider.setMaximum(100)
        self.security_level_slider.setPageStep(50)
        self.security_level_layout.addWidget(self.security_level_slider)

        self.left_layout.addLayout(self.security_level_layout)

        self.horizontalLayout_7 = QtWidgets.QHBoxLayout()
        self.horizontalLayout_7.setContentsMargins(0, 65, 0, 0)
        self.horizontalLayout_7.setObjectName("horizontalLayout_7")
        self.save_user_push = Push(self.horizontalLayout_7, 35, 35, 5,
                                   icon_path=os.path.join(CONTENT_PATH, 'media', 'buttons', 'save.svg'))
        self.save_user_push.setObjectName("save_user_push")
        self.horizontalLayout_7.addWidget(self.save_user_push)
        self.set_password_push = Push(self.horizontalLayout_7, 35, 35, 5,
                                      icon_path=os.path.join(CONTENT_PATH, 'media', 'buttons', 'set_password.svg'))
        self.set_password_push.setObjectName("set_password_push")
        self.horizontalLayout_7.addWidget(self.set_password_push)
        self.del_user_push = Push(self.horizontalLayout_7, 35, 35, 5,
                                  icon_path=os.path.join(CONTENT_PATH, 'media', 'buttons', 'del_user.svg'))
        self.del_user_push.setObjectName("del_user_push")
        self.horizontalLayout_7.addWidget(self.del_user_push)
        self.logout_push = Push(self.horizontalLayout_7, 35, 35, 5,
                                icon_path=os.path.join(CONTENT_PATH, 'media', 'buttons', 'logout.svg'))
        self.logout_push.setObjectName("logout_push")
        self.horizontalLayout_7.addWidget(self.logout_push)

        self.left_layout.addLayout(self.horizontalLayout_7)
        self.left_layout.addItem(QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Maximum,
                                                       QtWidgets.QSizePolicy.Policy.Expanding))

        self.center_layout.addLayout(self.left_layout)
        self.center_layout.addItem(QtWidgets.QSpacerItem(5000, 0, QtWidgets.QSizePolicy.Policy.Maximum,
                                                         QtWidgets.QSizePolicy.Policy.Expanding))

        self.verticalLayout_7 = QtWidgets.QVBoxLayout()
        self.verticalLayout_7.setObjectName("verticalLayout_7")
        self.verticalLayout_12 = QtWidgets.QVBoxLayout()
        self.verticalLayout_12.setObjectName("verticalLayout_12")
        self.photo = QtSvgWidgets.QSvgWidget(os.path.join(CONTENT_PATH, 'media', 'profile', f'{str(random.randint(1, 15))}.svg'))
        self.photo.setFixedSize(QtCore.QSize(250, 250))
        self.photo.setObjectName("photo")

        self.verticalLayout_12.addItem(
            QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Expanding))
        self.verticalLayout_12.addWidget(self.photo, QtCore.Qt.AlignmentFlag.AlignVCenter,
                                         QtCore.Qt.AlignmentFlag.AlignHCenter)

        self.verticalLayout_12.addItem(
            QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum, QtWidgets.QSizePolicy.Policy.Expanding))
        self.verticalLayout_7.addLayout(self.verticalLayout_12)

        self.verticalLayout_13 = QtWidgets.QVBoxLayout()
        self.verticalLayout_13.setObjectName("verticalLayout_13")

        self.verticalLayout_7.addLayout(self.verticalLayout_13)
        self.verticalLayout_7.addItem(QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Maximum,
                                                            QtWidgets.QSizePolicy.Policy.Expanding))
        self.center_layout.addLayout(self.verticalLayout_7)
        self.center_layout.addItem(QtWidgets.QSpacerItem(5000, 0, QtWidgets.QSizePolicy.Policy.Maximum,
                                                         QtWidgets.QSizePolicy.Policy.Expanding))

        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.base_layout.addWidget(self.scrollArea)

        self.setObjectName("profile")

        self.add_function()
        self.retranslateUi()

        self.is_exit = 0
        self.is_clicked = 0

    def load_atchivments(self):
        self.ACHIEVEMENT = {
            (0, 0): (self.tr('Поехали!'), os.path.join(CONTENT_PATH, 'media', 'achievements', '10.svg')),
            (0, 1): (self.tr('Полный набор'), os.path.join(CONTENT_PATH, 'media', 'achievements', '1.svg')),
            (0, 2): (self.tr('Игра окончена'), os.path.join(CONTENT_PATH, 'media', 'achievements', '2.svg')),
            (0, 3): (self.tr('Шерлок?'), os.path.join(CONTENT_PATH, 'media', 'achievements', '4.svg')),
            (0, 4): (self.tr('Исследователь'), os.path.join(CONTENT_PATH, 'media', 'achievements', '5.svg')),
            (1, 0): (self.tr('Делу время....'), os.path.join(CONTENT_PATH, 'media', 'achievements', '7.svg')),
            (1, 1): (self.tr('..потехе час'), os.path.join(CONTENT_PATH, 'media', 'achievements', '8.svg')),
            (1, 2): (self.tr('Частичка истории'), os.path.join(CONTENT_PATH, 'media', 'achievements', '14.svg')),
            (1, 3): (self.tr('Это что? Микробы!'), os.path.join(CONTENT_PATH, 'media', 'achievements', '12.svg')),
            (1, 4): (self.tr('Я видел темную сторону Луны'), os.path.join(CONTENT_PATH, 'media', 'achievements', '13.svg')),
            (2, 0): (self.tr('Моя прелесть'), os.path.join(CONTENT_PATH, 'media', 'achievements', '16.svg')),
            (2, 1): (self.tr('Сова'), os.path.join(CONTENT_PATH, 'media', 'achievements', '17.svg')),
            (2, 2): (self.tr('Классику сэр?'), os.path.join(CONTENT_PATH, 'media', 'achievements', '15.svg')),
            (2, 3): (self.tr('Трансформация'), os.path.join(CONTENT_PATH, 'media', 'achievements', '3.svg')),
            (2, 4): (self.tr('Быть или не быть'), os.path.join(CONTENT_PATH, 'media', 'achievements', '9.svg')),
            (3, 0): (self.tr('Ушел по-английски'), os.path.join(CONTENT_PATH, 'media', 'achievements', '11.svg')),
            (3, 1): (self.tr('Назад в будущее'), os.path.join(CONTENT_PATH, 'media', 'achievements', '18.svg')),
            (3, 2): (self.tr('Коллекционер'), os.path.join(CONTENT_PATH, 'media', 'achievements', '6.svg')),
            (3, 3): (self.tr('С НГ!'), os.path.join(CONTENT_PATH, 'media', 'achievements', '19.svg')),
            (3, 4): (self.tr('Звезда'), os.path.join(CONTENT_PATH, 'media', 'achievements', '20.svg')),
        }

    def add_function(self):
        self.save_user_push.clicked.connect(self.save_user)
        self.set_password_push.clicked.connect(self.set_password)
        self.del_user_push.clicked.connect(self.del_user)
        self.logout_push.clicked.connect(self.logout)
        self.photo.mousePressEvent = self.click_profile
        self.security_level_slider.sliderReleased.connect(self.click_slider)
        self.security_level_slider.valueChanged.connect(self.chaged_slider)
        self.security_level_slider.keyPressEvent = self.slider_key_pressEvent

    def retranslateUi(self):
        _translate = QtCore.QCoreApplication.translate
        self.accaunt_label.setText(self.tr("Профиль"))
        self.username_label.setText(self.tr("Имя пользователя"))
        self.fio_user_label.setText(self.tr("ФИО Своё"))
        self.teamleader_label.setText(self.tr("ФИО Кл.руководителя"))
        self.group_label.setText(self.tr("Группа"))
        self.specialization_label.setText(self.tr("Специальность"))

        self.security_leve_label.setText(self.tr('Уровень защиты данных'))

        self.save_user_push.setToolTip(self.tr('Сохранить изменения'))
        self.set_password_push.setToolTip(self.tr('Изменить пароль'))
        self.del_user_push.setToolTip(self.tr('Удалить пользователя'))
        self.logout_push.setToolTip(self.tr('Выйти из аккаунта'))
        self.load_atchivments()

        self.accaunt_label.setFont(QtGui.QFont(CM.NAME_FONT, 16 + CM.ADD_FONT_SIZE))
        self.username_label.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.fio_user_label.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.teamleader_label.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.group_label.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.specialization_label.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))

        self.security_leve_label.setFont(QtGui.QFont(CM.NAME_FONT, 16 + CM.ADD_FONT_SIZE))

    def click_profile(self, *args, **kwargs):
        photo = QtSvgWidgets.QSvgWidget(os.path.join(CONTENT_PATH, 'media', 'profile', f'{str(random.randint(1, 15))}.svg'))
        photo.setFixedSize(QtCore.QSize(250, 250))
        self.verticalLayout_12.replaceWidget(self.photo, photo)
        self.photo.deleteLater()
        self.photo = photo
        self.photo.mousePressEvent = self.click_profile
        if not self.is_clicked:
            self.cod = [2, 3]
            self.is_clicked = 1

    def init_atchivments(self):
        if hasattr(self, 'atchivmenys_layout'):
            self.verticalLayout_13 = QtWidgets.QHBoxLayout()
        if not USER_MANAGER.user.parametrs.get('achievements'):
            achievements = []
        else:
            achievements = USER_MANAGER.user.parametrs.get('achievements')

        for i in range(4):

            self.atchivmenys_layout = QtWidgets.QHBoxLayout()
            for j in range(5):
                if [i, j] in achievements:
                    qlable = QtSvgWidgets.QSvgWidget(self.ACHIEVEMENT[(i, j)][1])
                    qlable.setToolTip(str(self.ACHIEVEMENT[(i, j)][0]))
                    qlable.setFixedSize(QtCore.QSize(80, 80))
                else:
                    qlable = QtSvgWidgets.QSvgWidget(os.path.join(CONTENT_PATH, 'media', 'achievements', '0.svg'))
                    # qlable.setToolTip(str(self.ACHIEVEMENT[(i, j)][0]))
                    qlable.setFixedSize(QtCore.QSize(80, 80))
                self.atchivmenys_layout.addWidget(qlable)
            self.verticalLayout_13.addLayout(self.atchivmenys_layout)

    def __setattr__(self, key, value):
        if key == 'cod':
            if self.checking_condition(tuple(value)):
                user_atch = USER_MANAGER.user.parametrs.get('achievements')

                if user_atch:
                    USER_MANAGER.user.parametrs['achievements'].append(list(value))
                else:
                    USER_MANAGER.user.parametrs['achievements'] = [list(value)]

                USER_MANAGER.user.save_user()

        else:
            super().__setattr__(key, value)

    def save_user(self):
        try:
            if self.fio_user_edit.text():
                if USER_MANAGER.USER_CLASS.is_valid_fio(self.fio_user_edit.text()):
                    USER_MANAGER.user.parametrs['offical_name'] = self.fio_user_edit.text().title()
                else:
                    raise ValueError(self.tr(
                        'ФИО должно состоять только из букв и быть из 3 частей, каждая из которых не менее 2 символов'))
                self.fio_user_edit.setText(USER_MANAGER.user.parametrs.get('offical_name'))
            if self.fio_teamleader_edit.text():
                if USER_MANAGER.USER_CLASS.is_valid_fio(self.fio_teamleader_edit.text()):
                    USER_MANAGER.user.parametrs['teamleader'] = self.fio_teamleader_edit.text().title()
                else:
                    raise ValueError(self.tr(
                        'ФИО должно состоять только из букв и быть из 3 частей, каждая из которых не менее 2 символов'))
                self.fio_teamleader_edit.setText(USER_MANAGER.user.parametrs.get('teamleader'))
            USER_MANAGER.user.parametrs['group'] = self.group_edit.text()
            USER_MANAGER.user.parametrs['specialization'] = self.specialization_edit.text()
            if USER_MANAGER.user.parametrs['specialization'].lower() == 'великий сыщик':
                self.cod = [0, 3]



        except BaseException as f:
            self.message_profile.setText(str(f))
        else:
            self.message_profile.setText('')
            self.parent.get_down_message(self.tr('Данные профиля сохранены'))
            USER_MANAGER.user.save_user()

    def del_user(self):
        message = QtWidgets.QMessageBox.question(self.parent, self.tr('Удаление пользователя'),
                                                 self.tr("Вы точно хотите удалить свой аккаунт?"),
                                                 QtWidgets.QMessageBox.StandardButton.Yes | QtWidgets.QMessageBox.StandardButton.No)
        if message == message.Yes:
            USER_MANAGER.del_user()
            self.parent.status = 2
            self.parent.close()

    def logout(self):
        self.is_exit = 1
        self.cod = [3, 0]
        USER_MANAGER.user = None
        self.parent.status = 4
        self.parent.close()

    def set_password(self):
        window_password = WindowSetPassword(self.parent)
        window_password.retranslateUi()
        window_password.show()

    def update_user_info(self):
        self.username_edit.setText(USER_MANAGER.user.username)
        self.fio_teamleader_edit.setText(USER_MANAGER.user.parametrs.get('teamleader', ''))
        self.fio_user_edit.setText(USER_MANAGER.user.parametrs.get('offical_name', ''))
        self.group_edit.setText(USER_MANAGER.user.parametrs.get('group', ''))
        self.specialization_edit.setText(USER_MANAGER.user.parametrs.get('specialization', ''))

    def user_information_load(self):
        if not USER_MANAGER.user is None:
            self.update_photo_security_level(USER_MANAGER.user.security_level)
            self.security_level_slider.setValue(USER_MANAGER.user.security_level * 50)
            self.update_user_info()
            self.init_atchivments()

    def checking_condition(self, key):
        CONDITION = {
            (0, 0): 'True',
            (0, 1): 'len(MANAGER_STUDENTS.students) >= 25',
            (0, 2): 'self.parent.F6.is_click_end_period',
            (3, 2): 'len(self.parent.archive.files_archive) >= 12',
            (1, 0): 'self.parent.settings.is_add_work_day',
            (1, 1): 'self.parent.settings.is_del_work_day',
            (3, 0): 'self.is_exit',
            (1, 2): 'MANAGER_STUDENTS.period[1] == 1961',
            (1, 3): 'True',
            (1, 4): 'True',
            (3, 3): 'MANAGER_STUDENTS.period[0] == 12',
            (3, 1): 'self.parent.archive.is_comebake',
            (2, 1): 'datetime.datetime.today().time().hour >= 21',
            (2, 2): 'True',
            (2, 0): 'is_click_license',
            (2, 3): 'True',
            (2, 4): 'self.parent.archive.is_delite_file',
            (0, 4): 'len(self.parent.tab_set) >= 5',
            (0, 3): 'True',
            (3, 4): 'len(USER_MANAGER.user.parametrs.get("achievements")) >= 19',

        }
        if CONDITION.get(tuple(key)):
            if USER_MANAGER.user.parametrs:
                if USER_MANAGER.user.parametrs.get('achievements'):
                    if not list(key) in USER_MANAGER.user.parametrs.get('achievements'):
                        return eval(CONDITION[tuple(key)])

    def click_slider(self):
        value = self.security_level_slider.value()
        if not value in [0, 50, 100]:
            if value > 50:
                if value > 75:
                    self.security_level_slider.setValue(100)
                else:
                    self.security_level_slider.setValue(50)
                self.click_slider()
            else:
                if value > 25:
                    self.security_level_slider.setValue(50)
                else:
                    self.security_level_slider.setValue(0)
                self.click_slider()

    def slider_key_pressEvent(self, event):
        if event.key() == QtCore.Qt.Key.Key_Right:
            if self.security_level_slider.value() + 50 <= 100:
                self.security_level_slider.setValue(self.security_level_slider.value() + 50)
        elif event.key() == QtCore.Qt.Key.Key_Left:
            if self.security_level_slider.value() - 50 >= 0:
                self.security_level_slider.setValue(self.security_level_slider.value() - 50)
        else:
            QtWidgets.QWidget.keyPressEvent(self, event)

    def chaged_slider(self):
        value = self.security_level_slider.value()
        if value in [0, 50, 100]:
            self.security_level_photo.deleteLater()
            self.security_level_photo = QtSvgWidgets.QSvgWidget(
                os.path.join(CONTENT_PATH, 'media', 'security_level', f'security_level_{str(value // 50)}.svg'))
            self.security_level_photo.setFixedSize(100, 100)

            self.security_level_photo_layout.addWidget(self.security_level_photo, QtCore.Qt.AlignmentFlag.AlignVCenter,
                                                       QtCore.Qt.AlignmentFlag.AlignHCenter)

            if value // 50 != MANAGER_STUDENTS.user.security_level:
                MANAGER_STUDENTS.user.security_level = value // 50
                USER_MANAGER.user.save_user()
                MANAGER_STUDENTS.save_students()

    def update_photo_security_level(self, level=None):
        if hasattr(self, 'security_level_photo'):
            self.security_level_photo.deleteLater()
        if isinstance(level, int) and 0 <= level <= 2:
            self.security_level_photo = QtSvgWidgets.QSvgWidget(
                os.path.join(CONTENT_PATH, 'media', 'security_level', f'security_level_{str(level)}.svg'))
        else:
            self.security_level_photo = QtSvgWidgets.QSvgWidget(
                os.path.join(CONTENT_PATH, 'media', 'security_level', 'security_level_none.svg'))

        self.security_level_photo.setFixedSize(100, 100)
        self.security_level_photo_layout.addWidget(self.security_level_photo, QtCore.Qt.AlignmentFlag.AlignVCenter,
                                                   QtCore.Qt.AlignmentFlag.AlignHCenter)


class MplCanvas(FigureCanvas):

    def __init__(self, parent=None, width=5, height=4, dpi=100):
        fig = Figure(figsize=(width, height), dpi=dpi)
        self.axes = fig.add_subplot()
        super(MplCanvas, self).__init__(fig)


class StatisticTab(QtWidgets.QWidget):
    STYLE = ['Solarize_Light2', '_classic_test_patch', '_mpl-gallery', '_mpl-gallery-nogrid', 'bmh', 'seaborn-v0_8',
             'seaborn-v0_8-bright',
             'seaborn-v0_8-colorblind', 'seaborn-v0_8-dark', 'seaborn-v0_8-dark-palette', 'seaborn-v0_8-darkgrid',
             'seaborn-v0_8-deep', 'seaborn-v0_8-muted', 'seaborn-v0_8-notebook', 'seaborn-v0_8-paper',
             'seaborn-v0_8-pastel', 'seaborn-v0_8-talk', 'seaborn-v0_8-whitegrid', 'tableau-colorblind10']

    def __init__(self, parent=None):
        super(StatisticTab, self).__init__()
        self.parent = parent
        self.setObjectName(u"tab_2")
        self.setStyleSheet("""
                #button_before {
                background-color: green;
                font-size: 4em;
                border: none;
                color: white;
            }""")

        self.verticalLayout_14 = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_14.setObjectName(u"verticalLayout_14")

        self.button_before = QtWidgets.QPushButton()
        self.button_before.setObjectName('button_before')
        font = QtGui.QFont()
        font.setPointSize(14)
        font.setBold(False)
        font.setItalic(False)
        font.setUnderline(False)
        font.setWeight(QtGui.QFont.Weight(50))
        self.button_before.setFont(font)
        self.styles_diagram = ManagerStudents.crate_eternal_iter(self.STYLE)

        self.button_before.setFixedSize(175, 70)
        self.button_before.clicked.connect(self.clicked_before)
        self.verticalLayout_14.addWidget(self.button_before, QtCore.Qt.AlignmentFlag.AlignVCenter,
                                         QtCore.Qt.AlignmentFlag.AlignHCenter)

    def retranslateUi(self):
        self.button_before.setText(self.tr("Создать статистику"))

        if hasattr(self, 'scrollArea_2'):
            self.retranslate()

    def init_statistic(self):
        self.setStyleSheet(
            """
            #get_month_statistic {
               font-size: 16px;
               color: black;
               background-color: Black;
               border: none;
               padding: 5px;
               color: white;
               margin: 0px;
               
               }

            
               
            #get_more_statistic {
               font-size: 16px;
               color: black;
               background-color: Black;
               border: none;
               padding: 5px;
               color: white;
               margin: 0px;
               
               }
            
            #scrollAreaWidgetContents_2 {background-color:white; background-color:white}
            QTableWidget {border: none; background-color:white;}
         
        )"""
        )
        self.scrollArea_2 = QtWidgets.QScrollArea()
        self.scrollArea_2.setObjectName(u"scrollArea_2")
        self.scrollArea_2.setWidgetResizable(True)
        self.scrollAreaWidgetContents_2 = QtWidgets.QWidget()
        self.scrollAreaWidgetContents_2.setObjectName(u"scrollAreaWidgetContents_2")

        self.verticalLayout_0 = QtWidgets.QVBoxLayout(self.scrollAreaWidgetContents_2)

        self.horizontalLayout_buttons_up = QtWidgets.QHBoxLayout()

        self.get_month_statistic = QtWidgets.QPushButton()
        self.get_month_statistic.setObjectName(u'get_month_statistic')

        self.horizontalLayout_buttons_up.addWidget(self.get_month_statistic)

        self.get_more_statistic = QtWidgets.QPushButton()

        self.get_more_statistic.setObjectName(u'get_more_statistic')
        self.horizontalLayout_buttons_up.addWidget(self.get_more_statistic)

        self.verticalLayout_0.addLayout(self.horizontalLayout_buttons_up)
        self.message_none = QtWidgets.QLabel()
        self.message_none.hide()
        self.message_none.setStyleSheet('color: red;')
        self.message_none.setAlignment(QtCore.Qt.AlignCenter)
        self.message_none.setFont(QtGui.QFont(CM.NAME_FONT, CM.ADD_FONT_SIZE + 14))
        self.verticalLayout_0.addWidget(self.message_none)
        # self.verticalLayout_16 = QtWidgets.QVBoxLayout()
        # self.verticalLayout_16.setObjectName(u"verticalLayout_16")
        # # self.grafic_group = Grafics()
        # # self.grafic_group.draw_circle_graph(is_number=False)

        self.canvas = MplCanvas(self, width=15, height=5, dpi=100)
        self.clicked_get_month_statistic()

        # Setup a timer to trigger the redraw by calling update_plot.
        # self.timer = QTimer()
        # self.timer.setInterval(100)
        # self.timer.timeout.connect(self.update_plot)
        # self.timer.start()
        # self.timer = QTimer()
        # self.timer.setInterval(100)
        # self.timer.timeout.connect(self.update_plot)
        # self.timer.start()

        # self.verticalLayout_16.addWidget(self.canvas)

        self.add_function()
        self.verticalLayout_0.addWidget(self.canvas)

        self.retranslate()
        self.scrollArea_2.setWidget(self.scrollAreaWidgetContents_2)
        self.verticalLayout_14.addWidget(self.scrollArea_2)

    def retranslate(self):
        self.message_none.setText(self.tr('Недостаточно данных для отображения графика'))
        self.get_month_statistic.setText(self.tr('Статистика за месяц'))
        self.get_more_statistic.setText(self.tr('Статистика за месяцы'))

    def update_plot(self, is_circle=True):
        self.canvas.axes.cla()
        is_draw = 1
        if is_circle:
            stat = MANAGER_STUDENTS.get_statistics_for_group()
            if stat:
                self.xdata = list(map(lambda x: MANAGER_STUDENTS.CLASS_STUDENT.create_shorts_fio(x) + f'({stat[x]!r})',
                                      list(stat.keys())))
                self.ydata = list(stat.values())
                self.canvas.axes.pie(self.ydata, labels=self.xdata, autopct='%1.2f%%')
                self.canvas.axes.set_title(self.tr('Доля прогулов на студента'))
            else:
                is_draw = 0

        else:
            if CM.CURRENT_LANGUAGE:
                if CM.CURRENT_LANGUAGE == 'china':
                    font = font_manager.FontProperties(family='Microsoft JhengHei',
                                                       weight='bold',
                                                       style='normal', size=12 + CM.ADD_FONT_SIZE)
                else:
                    font = font_manager.FontProperties(family='Arial',
                                                       weight='bold',
                                                       style='normal', size=12 + CM.ADD_FONT_SIZE)

            statistic = MANAGER_STUDENTS.get_total_statistic_period()
            if len(statistic) > 1:
                labels = statistic.keys()
                labels = list(map(lambda x: f'{x[0]!r} {ManagerStudents.MONTHS[x[1]]}', labels))
                data = list(map(lambda x: x[0], statistic.values()))
                data1 = list(map(lambda x: x[1], statistic.values()))
                data2 = list(map(lambda x: x[2], statistic.values()))
                xs = [len(labels[0]) * 20 * i for i in range(len(labels))]
                self.canvas.axes.plot(xs, data,
                                      label=self.tr('ПОУВ'))
                self.canvas.axes.plot(xs, data1, label=self.tr('НЕУВ'))
                self.canvas.axes.plot(xs, data2, label=self.tr('Всего'))

                for i, v in zip(xs, data):
                    self.canvas.axes.text(i, v + 25, "%d" % v, ha="center")
                for i, v in zip(xs, data1):
                    self.canvas.axes.text(i, v + 25, "%d" % v, ha="center")
                for i, v in zip(xs, data2):
                    self.canvas.axes.text(i, v + 25, "%d" % v, ha="center")

                self.canvas.axes.set_xticks([len(labels[0]) * 20 * i for i in range(len(labels))], labels, rotation=45)
                self.canvas.axes.legend(prop=font)
                self.canvas.axes.set_ylabel(self.tr('Количество прогулов'))
                self.canvas.axes.set_xlabel(self.tr('Месяцы'))
                self.canvas.axes.set_title(self.tr('Тенденция прогулов'))

            else:
                is_draw = 0

        if is_draw:
            self.message_none.hide()
            self.canvas.draw()
        else:
            self.message_none.show()

    def add_function(self):
        self.get_month_statistic.clicked.connect(self.clicked_get_month_statistic)
        self.get_more_statistic.clicked.connect(self.clicked_get_period_statistics)

    def clicked_get_month_statistic(self):
        self.get_month_statistic.setStyleSheet('color: black; background-color: white; border: 2px solid black')
        self.get_more_statistic.setStyleSheet('color: white; background-color: black;')
        self.update_plot(True)

    def clicked_get_period_statistics(self):
        self.get_more_statistic.setStyleSheet('color: black; background-color: white; border: 2px solid black')
        self.get_month_statistic.setStyleSheet('color: white; background-color: black;')
        self.update_plot(False)

    def clicked_before(self):
        self.init_statistic()
        self.button_before.hide()


class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        self.status = 0
        super().__init__()
        self.setStyleSheet(
            """
            #game_over_push, #save_to_exel_push, #save_table_push {border: none;}
            #QButton {margin: auto;}
            QTableWidget {margin-left: auto; margin-right: auto;}
            QListWidget {margin-left: auto; margin-top: 20%; margin-right: auto; border: none;}
            #profile QPushButton {background-color: rgba(249, 248, 244, 0);}
            QScrollArea {border: none}
    
            

            """)
        self.setWindowIcon(QtGui.QIcon(os.path.join(CONTENT_PATH, 'media', 'logo.svg')))
        self.setObjectName("MainWindow")
        self.resize(1024, 768)
        self.setMinimumSize(800, 450)
        self.centralwidget = QtWidgets.QWidget(self)
        self.centralwidget.setObjectName("centralwidget")
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self.centralwidget)
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.group = QtWidgets.QTabWidget(self.centralwidget)
        self.group.setTabPosition(QtWidgets.QTabWidget.TabPosition.North)
        self.group.setTabShape(QtWidgets.QTabWidget.TabShape.Rounded)
        self.group.setDocumentMode(False)
        self.group.setTabsClosable(False)
        self.group.setMovable(False)
        self.group.setTabBarAutoHide(False)
        self.group.setObjectName("group")

        self.F6 = AbsenceTab(self)
        self.group.addTab(self.F6, "")

        self.marks = MarksTab(self)
        self.group.addTab(self.marks, "")

        self.students = StudentsTab(self)
        self.group.addTab(self.students, "")

        self.archive = ArchiveTab(self)
        self.group.addTab(self.archive, "")

        self.settings = SettingsTab(self)
        self.group.addTab(self.settings, "")

        self.statistics = StatisticTab(self)
        self.group.addTab(self.statistics, "")

        self.profile = ProfileTab(self)
        self.group.addTab(self.profile, "")

        self.text_massage = QtWidgets.QLabel()
        self.text_massage.hide()
        self.text_massage.setFont(QtGui.QFont(CM.NAME_FONT, 12 + CM.ADD_FONT_SIZE))
        self.text_massage.setAlignment(QtCore.Qt.AlignmentFlag.AlignRight)
        self.verticalLayout_2.addWidget(self.group)

        self.verticalLayout_2.addWidget(self.text_massage)
        self.setCentralWidget(self.centralwidget)

        self.retranslateUi()
        self.group.setCurrentIndex(0)
        QtCore.QMetaObject.connectSlotsByName(self)

        self.tab_set = set()
        self.add_function()

        update_massage_dialog = UpdateManager(self)
        update_massage_dialog.show()

    def retranslateUi(self):

        if hasattr(self, 'marks'):
            self.group.setTabText(self.group.indexOf(self.marks), self.tr("Оценки"))
        self.setWindowTitle(self.tr("F6"))

        self.group.setTabText(self.group.indexOf(self.F6), self.tr("Прогулы"))
        self.group.setTabText(self.group.indexOf(self.students), self.tr("Студенты"))
        self.group.setTabText(self.group.indexOf(self.settings), self.tr("Настройки"))
        self.group.setTabText(self.group.indexOf(self.profile), self.tr("Профиль"))
        self.group.setTabText(self.group.indexOf(self.archive), self.tr("Архив"))
        self.group.setTabText(self.group.indexOf(self.statistics), self.tr("Статистика"))

        self.F6.retranslateUi()

        self.marks.retranslateUi()

        self.students.retranslateUi()
        self.profile.retranslateUi()
        self.archive.retranslateUi()
        self.settings.retranslateUi()
        self.statistics.retranslateUi()

        self.group.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        CM.IS_CHANGE = False

    def keyPressEvent(self, e):
        k = e.key()
        super().keyPressEvent(e)
        if k == 83:
            if not self.only_show:
                self.F6.save_table()
        CM.IS_CHANGE = False

    def get_down_message(self, message, is_error=False, time=2000):
        self.text_massage.setText(message)
        if not is_error:
            self.text_massage.setStyleSheet('color: green')
        else:
            self.text_massage.setStyleSheet('color: red')
        self.text_massage.show()
        self.timer = QtCore.QTimer()
        self.timer.timeout.connect(lambda: self.text_massage.hide())
        self.timer.start(time)

    def click_tab(self):
        if self.group.currentIndex() not in self.tab_set:
            self.tab_set.add(self.group.currentIndex())

        self.profile.cod = [0, 4]

    def add_function(self):
        pass

    def init_students_manager(self, path=None, only_show=False, period=None):
        self.only_show = only_show

        self.profile.cod = [3, 4]
        self.profile.cod = [2, 0]

        global MANAGER_STUDENTS
        if not [0, 0] in USER_MANAGER.user.parametrs.get('achievements', []):
            USER_MANAGER.user.add_achievement([0, 0])
        if not period:
            try:
                MANAGER_STUDENTS = ManagerStudents.load_manager_students(USER_MANAGER.user, file_name=path)
            except BaseException as message:
                if 'archive' in path:
                    raise FileNotFoundError

                MANAGER_STUDENTS = ManagerStudents((datetime.date.today().month, datetime.date.today().year),
                                                   USER_MANAGER.user)


        else:
            students = []
            if not MANAGER_STUDENTS is None:
                students = MANAGER_STUDENTS.students
                for i in range(len(students)):
                    students[i].sick_days.clear()
                    students[i].absence_days.clear()
                    students[i].marks.clear()
            MANAGER_STUDENTS = ManagerStudents(period,
                                               USER_MANAGER.user, students=students)
        if len(MANAGER_STUDENTS.students) == 0:
            self.group.removeTab(self.group.indexOf(self.students))

        try:
            self.F6.init_table_absence(only_show)
            if hasattr(self, 'marks'):
                self.marks.init_table_marks(only_show)
        except BaseException as f:
            print('Ошибка инициализации таблиц', repr(f))

        self.students.update_list_students()

        self.archive.init_archive()

        self.profile.user_information_load()

        if [0, 4] not in (USER_MANAGER.user.parametrs.get('achievements') or []):
            self.group.tabBarClicked.connect(self.click_tab)

        if USER_MANAGER.user.parametrs:
            if USER_MANAGER.user.parametrs.get('table_marks') == False:
                self.group.removeTab(self.group.indexOf(self.marks))

            if USER_MANAGER.user.parametrs.get('statistic') == False:
                self.group.removeTab(self.group.indexOf(self.statistics))

        if len(self.archive.files_archive) == 0 and self.group.indexOf(self.archive) != -1:
            self.group.removeTab(self.group.indexOf(self.archive))

        self.retranslateUi()

        self.profile.cod = [1, 2]
        self.profile.cod = [3, 3]
        self.profile.cod = [2, 1]
        self.profile.cod = [3, 4]
        global is_click_license
        is_click_license = 0
        self.group.currentIndex()


class BaseTable(QtWidgets.QTableWidget):
    def __init__(self, only_show=False):
        super().__init__()
        self.setStyleSheet("""QTableWidget {border:red;}""")
        self.only_show = only_show

    def update_table_students(self, *args, **kwargs):
        """Очищает полностью таблицу и создает ее заново"""
        self.clear()
        self.generate_table(*args, **kwargs)
        if len(MANAGER_STUDENTS.students) >= 25:
            if not [0, 1] in USER_MANAGER.user.parametrs.get('achievements', []):
                USER_MANAGER.user.add_achievement([0, 1])


class TableAbsence(BaseTable, QtWidgets.QTableWidget):
    """Модуль TableAbsence отвечает за создание и обновление таблицы прогулов принимает на вход режим отображения:
    атрибут only_show, который может быть логическим(True/False)"""

    def retranslateUi(self, size=0):
        self.setItem(0, 0, QTableWidgetItem(
            self.tr(
                'ВЕДОМОСТЬ УЧЁТА ЧАСОВ ПРОГУЛОВ за') + f" {str(ManagerStudents.MONTHS[MANAGER_STUDENTS.period[0] - 1]) + ' ' + str(MANAGER_STUDENTS.period[1]) if CM.CURRENT_LANGUAGE == 'russia' else str((MANAGER_STUDENTS.period[0], MANAGER_STUDENTS.period[1]))}"))
        title = self.item(0, 0)
        title.setBackground(QtGui.QColor(153, 153, 153))
        title.setFont(QtGui.QFont('Calibri', 26 + size))
        title.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)

        self.setItem(1, 34, QTableWidgetItem(self.tr("Из них")))
        self.item(1, 34).setFont(QtGui.QFont('Calibri', 14 + size))
        self.item(1, 34).setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
        self.item(1, 34).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)

        self.setItem(2, 34, QTableWidgetItem(self.tr("УВ")))
        self.item(2, 34).setFont(QtGui.QFont('Calibri', 14 + size))
        self.item(2, 34).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.item(2, 34).setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)

        self.setItem(2, 35, QTableWidgetItem(self.tr("НЕУВ")))
        self.item(2, 35).setFont(QtGui.QFont('Calibri', 14 + size))
        self.item(2, 35).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.item(2, 35).setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)

        self.setItem(1, 1, QTableWidgetItem(self.tr("ФИО")))
        fio = self.item(1, 1)
        fio.setFont(QtGui.QFont('Calibri', 14 + size))
        fio.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        fio.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
        self.setItem(1, 33, QTableWidgetItem(self.tr("Итог")))
        result_up = self.item(1, 33)
        result_up.setFont(QtGui.QFont('Calibri', 14 + size))
        result_up.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)

    def generate_table(self, size=0):
        self.mod_size = size
        self.setColumnCount(36)

        if not self.only_show and len(MANAGER_STUDENTS.students) < 30:
            self.setRowCount(2 + len(MANAGER_STUDENTS.students) + 2)
        else:
            self.setRowCount(2 + len(MANAGER_STUDENTS.students) + 1)

        self.horizontalHeader().setVisible(False)
        self.verticalHeader().setVisible(False)
        self.setSpan(0, 0, 1, 36)
        self.setSpan(1, 1, 2, 1)
        self.setSpan(1, 34, 1, 2)
        self.setSpan(1, 0, 2, 1)

        self.retranslateUi(size)

        self.setItem(2, 33, QTableWidgetItem(str(sum(MANAGER_STUDENTS.days.values()))))

        self.update_hours_by_days(size)

        for i, student in enumerate(MANAGER_STUDENTS.students):
            i += 3
            self.setItem(i, 1, QTableWidgetItem(student.create_shorts_fio(student.fio)))
            self.item(i, 1).setFont(QtGui.QFont('Calibri', 14 + size))
            self.item(i, 1).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.item(i, 1).setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.setItem(i, 0, QTableWidgetItem(str(i - 2)))
            self.item(i, 0).setFont(QtGui.QFont('Calibri', 14 + size))

            self.setItem(i, 34, QTableWidgetItem(str(i - 2)))
            self.setItem(i, 35, QTableWidgetItem(str(i - 2)))

            for s_d in student.sick_days:
                self.setItem(i, s_d + 1, QTableWidgetItem(''))
                self.item(i, s_d + 1).setFont(QtGui.QFont('Calibri', 14 + size))
                self.add_hours_in_table(i, s_d + 1, student.sick_days.get(s_d), type_day='s')

            for a_d in student.absence_days:
                self.setItem(i, a_d + 1, QTableWidgetItem(''))
                self.item(i, a_d + 1).setFont(QtGui.QFont('Calibri', 14 + size))
                self.add_hours_in_table(i, a_d + 1, student.absence_days.get(a_d), type_day='a')
            self.update_statistics_student(i)
            self.update_hours_day()
            self.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.ResizeMode.Stretch)
            self.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.ResizeToContents)
            self.horizontalHeader().setSectionResizeMode(1, QtWidgets.QHeaderView.ResizeMode.ResizeToContents)
            self.horizontalHeader().setSectionResizeMode(33, QtWidgets.QHeaderView.ResizeMode.ResizeToContents)
            self.horizontalHeader().setSectionResizeMode(34, QtWidgets.QHeaderView.ResizeMode.ResizeToContents)
            self.horizontalHeader().setSectionResizeMode(35, QtWidgets.QHeaderView.ResizeMode.ResizeToContents)
            self.verticalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.ResizeToContents)

    def add_hours_in_table(self, row: int, column: int, hours: int, type_day: str = 's'):
        """Функция принимает координаты ячейки и число прогулов(часы) и
        тип дня(бывает a – прогулы по неуважительной причине и s – прогулы по болезни).
        Функция добавляет прогулы к объекту студент и в таблицу для отображения в приложении"""

        if type_day.lower() == 's':
            if column - 1 in MANAGER_STUDENTS.students[row - 3].absence_days:
                del MANAGER_STUDENTS.students[row - 3].absence_days[column - 1]
        elif type_day.lower() == 'a':
            if column - 1 in MANAGER_STUDENTS.students[row - 3].sick_days:
                del MANAGER_STUDENTS.students[row - 3].sick_days[column - 1]
        MANAGER_STUDENTS.add_day(
            MANAGER_STUDENTS.students[row - 3],
            column - 1,
            int(hours),
            type_day=type_day,
        )
        self.item(row, column).setText(str(hours))
        if type_day.lower() == 's':
            self.item(row, column).setBackground(QtGui.QColor(51, 153, 51))
        else:
            self.item(row, column).setBackground(QtGui.QColor(255, 165, 0))

        self.item(row, column).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.item(row, column).setFont(QtGui.QFont('Calibri', 14 + self.mod_size))

    def update_hours_day(self):
        """Обновляет значения часов для всех рабочих дней"""
        self.item(2, 33).setText(str(sum(MANAGER_STUDENTS.days.values())))
        self.item(2, 33).setFont(QtGui.QFont('Calibri', 14 + self.mod_size))

    def update_hours_by_days(self, size=0):
        for i in range(2, 32 + 1):
            if i - 1 in MANAGER_STUDENTS.days:
                self.setItem(1, i, QTableWidgetItem(
                    str(MANAGER_STUDENTS.days[i - 1] if MANAGER_STUDENTS.days[i - 1] else '')))
                self.item(1, i).setFont(QtGui.QFont('Calibri', 14 + size))
            else:
                self.setItem(1, i, QTableWidgetItem(str('✖')))
                self.item(1, i).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
                self.item(1, i).setFont(QtGui.QFont('Calibri', 14 + size))
                self.item(1, i).setBackground(QtGui.QColor(220, 220, 220))
                self.item(1, i).setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)

                try:
                    for j in range(2, self.rowCount()):
                        self.setItem(j, i, QTableWidgetItem(""))
                        self.item(j, i).setBackground(QtGui.QColor(220, 220, 220))

                        self.item(j, i).setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)

                except BaseException as f:
                    print(f)
            self.setItem(2, i, QTableWidgetItem(str(i - 1)))
            self.item(2, i).setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.item(2, i).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.item(2, i).setFont(QtGui.QFont('Calibri', 14 + size))
            if not i - 1 in MANAGER_STUDENTS.days:
                self.item(2, i).setBackground(QtGui.QColor(220, 220, 220))
            self.setColumnWidth(i, 10)

    def update_statistics_student(self, row):
        """Обновляет значения прогулов для всех студентов"""
        statistics = MANAGER_STUDENTS.students[row - 3].get_statistic_for_student()
        self.item(row, 34).setText(
            str(statistics['Sick_days'][1] if statistics['Sick_days'][1] > 0 else ''))
        self.item(row, 35).setText(
            str(statistics['Absence_days'][1] if statistics['Absence_days'][1] > 0 else ''))
        self.item(row, 1).setToolTip(
            self.tr('ФИО') + f": {statistics['FIO']};\n" +
            self.tr('Прогулы по неув. причине') + f': {str(statistics["Absence_days"][1])};\n' +
            self.tr('Прогулы по ув. причине') + f': {str(statistics["Sick_days"][1])}')
        self.item(row, 34).setFont(QtGui.QFont('Calibri', 14 + self.mod_size))
        self.item(row, 35).setFont(QtGui.QFont('Calibri', 14 + self.mod_size))


class TableMarks(BaseTable, QtWidgets.QTableWidget):
    """Модуль TableMarks отвечает за создание и обновление таблицы оценок, принимает на вход режим отображения:
        атрибут only_show, который может быть логическим(True/False)"""

    def retranslateUi(self, size=0):
        self.setItem(0, 0, QTableWidgetItem(
            self.tr(
                'ВЕДОМОСТЬ УЧЁТА УСПЕВАЕМОСТИ за ') + f" {str(ManagerStudents.MONTHS[MANAGER_STUDENTS.period[0] - 1]) + ' ' + str(MANAGER_STUDENTS.period[1]) if CM.CURRENT_LANGUAGE == 'russia' else str((MANAGER_STUDENTS.period[0], MANAGER_STUDENTS.period[1]))}"))
        title = self.item(0, 0)
        title.setBackground(QtGui.QColor(153, 153, 153))
        title.setFont(QtGui.QFont('Calibri', 26 + size))
        title.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)

        self.setItem(1, 1, QTableWidgetItem(self.tr("ФИО")))
        fio = self.item(1, 1)
        fio.setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        fio.setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
        fio.setFont(QtGui.QFont('Calibri', 14 + size))

    def generate_table(self, size=0):
        self.mod_size = size

        self.setColumnCount(13)
        if not self.only_show and len(MANAGER_STUDENTS.students) < 30:
            self.setRowCount(2 + len(MANAGER_STUDENTS.students) + 2)
        else:
            self.setRowCount(2 + len(MANAGER_STUDENTS.students) + 1)
        self.horizontalHeader().setVisible(False)
        self.verticalHeader().setVisible(False)
        self.setSpan(0, 0, 1, 36)
        self.setSpan(1, 1, 2, 1)
        self.setSpan(1, 34, 1, 2)
        self.setSpan(1, 0, 2, 1)

        self.retranslateUi(size)

        for i in range(2, 13):
            self.setItem(1, i, QTableWidgetItem(MANAGER_STUDENTS.couples.get(str(i - 1), '  ')[0]))
            self.setItem(2, i, QTableWidgetItem(MANAGER_STUDENTS.couples.get(str(i - 1), '  ')[1]))
            self.item(1, i).setFont(QtGui.QFont('Calibri', 14 + size))
            self.item(2, i).setFont(QtGui.QFont('Calibri', 14 + size))

        for i, student in enumerate(MANAGER_STUDENTS.students):
            i += 3
            self.setItem(i, 0, QTableWidgetItem(str(i - 2)))
            self.item(i, 0).setFont(QtGui.QFont('Calibri', 14 + size))
            self.setItem(i, 1, QTableWidgetItem(student.create_shorts_fio(student.fio)))
            self.item(i, 1).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
            self.item(i, 1).setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
            self.item(i, 1).setFont(QtGui.QFont('Calibri', 14 + size))
            self.update_statistics_student(i)
            for j in MANAGER_STUDENTS.students[i - 3].marks:
                self.setItem(i, j + 1, QTableWidgetItem(''))
                self.add_mark_table(i, j + 1, MANAGER_STUDENTS.students[i - 3].marks[j], size=size)
        # self.resizeColumnsToContents()
        self.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.ResizeMode.Stretch)
        self.horizontalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.ResizeToContents)
        self.horizontalHeader().setSectionResizeMode(1, QtWidgets.QHeaderView.ResizeMode.ResizeToContents)
        self.verticalHeader().setSectionResizeMode(0, QtWidgets.QHeaderView.ResizeMode.ResizeToContents)

    def add_mark_table(self, row, column, mark, size=0):
        mark = int(mark)
        MANAGER_STUDENTS.students[row - 3].marks[column - 1] = mark
        self.item(row, column).setText(str(mark))
        self.item(row, column).setFont(QtGui.QFont('Calibri', 14 + size))
        self.item(row, column).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        if mark == 5:
            self.item(row, column).setBackground(QtGui.QColor('#f1c30f'))
        elif mark == 4:
            self.item(row, column).setBackground(QtGui.QColor('#f39c11'))
        elif mark == 3:
            self.item(row, column).setBackground(QtGui.QColor('#d15400'))
        else:
            self.item(row, column).setBackground(QtGui.QColor('#c2392c'))

    def del_mark_table(self, row, column):
        if MANAGER_STUDENTS.students[row - 3].marks.get(column - 1):
            del MANAGER_STUDENTS.students[row - 3].marks[column - 1]
            self.item(row, column).setBackground(QtGui.QColor('#fff'))

    def update_statistics_student(self, row):
        statistics = MANAGER_STUDENTS.students[row - 3].get_statistic_for_student()
        self.item(row, 1).setToolTip(f"""
    {self.tr('ФИО')}: {statistics['FIO']};
                        """)

    def add_couples(self, number, couple=None, fio=None):
        if fio:
            fio = USER_MANAGER.user.check_fio(fio)
        # _______________________________________________________ ______________________

        if MANAGER_STUDENTS.couples.get(str(number)):
            row = MANAGER_STUDENTS.couples[str(number)]
            if couple:
                row[0] = couple if couple else ''
            if fio:
                row[1] = fio if fio else ''

            MANAGER_STUDENTS.couples[str(number)] = row
        else:
            MANAGER_STUDENTS.couples[str(number)] = [couple if couple else '', fio if fio else '']


class MonthCalendarView(QtWidgets.QDialog):
    def __init__(self, manager, period=None, parent=None, size=0, is_only_happy_days=False):
        super(MonthCalendarView, self).__init__(parent)
        self.setStyleSheet("* {border: none;}")
        self.setObjectName("ControlDays")
        self.setWindowTitle(self.tr("Месяц") + " " + str(manager.period)[1:-1])

        self.setModal(True)
        self.manager = manager
        if period:
            self.period = period
        else:
            self.period = self.manager.period
        self.parent = parent
        self.size = size
        self.is_only_happy_days = is_only_happy_days

        self.vertical_layout = QtWidgets.QVBoxLayout(self)

        if LANGUAGES == 'rassia':
            self.name_month = QtWidgets.QLabel(f'{str(manager.MONTHS[self.period[0] - 1])} {str(self.period[1])}')
        else:
            self.name_month = QtWidgets.QLabel(self.tr("Месяц") + " " + str(self.period)[1:-1])
        self.name_month.setFont(QtGui.QFont(CM.NAME_FONT, 16 + CM.ADD_FONT_SIZE))
        self.name_month.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.vertical_layout.addWidget(self.name_month)

        self.table_layout = QtWidgets.QVBoxLayout()

        self.update_table(self.is_only_happy_days)




    def click_cell(self, row, column):
        if self.period[0] == self.manager.period[0]:
            if self.table.item(row, column).text().isnumeric():
                if self.table.item(row, column).background() == QtGui.QColor(0, 0, 0, 0):
                    self.manager.off_day(int(self.table.item(row, column).text()))
                    self.table.item(row, column).setBackground(QtGui.QColor(60, 170, 60))
                    self.parent.F6.init_table_absence()
                elif self.table.item(row, column).background() == QtGui.QColor(60, 170, 60):
                    self.manager.on_day(int(self.table.item(row, column).text()))
                    self.table.item(row, column).setBackground(QtGui.QColor(0, 0, 0, 0))
                    self.parent.F6.init_table_absence()
        else:
            self.click_double_cell(row, column)

    def update_table(self, is_only_happy_days=False):
        if hasattr(self, 'table'):
            del self.table
            self.table_layout.deleteLater()

        self.table = QtWidgets.QTableWidget()
        if is_only_happy_days:
            self.table.setMinimumSize(350, 250)
        else:
            self.table.setMinimumSize(450, 250)

        self.table.setColumnCount(7)

        self.table.setRowCount(len(list(calendar.monthcalendar(month=self.period[0], year=self.period[1]))))
        row_id = 0
        if not is_only_happy_days:
            for week in calendar.monthcalendar(month=self.period[0], year=self.period[1]):
                for i in range(len(week)):
                    if not week[i]:
                        self.table.setItem(row_id, i, QTableWidgetItem(str('✖')))
                    elif (not (week[i] in self.manager.days)) and (self.manager.user.happy_days.get(str(self.period[0])) and
                                                              week[i] in self.manager.user.happy_days.get(
                                str(self.period[0]))):
                        self.table.setItem(row_id, i, QTableWidgetItem(str(week[i])))
                        self.table.item(row_id, i).setBackground(QtGui.QColor(181, 71, 71))
                    elif not (week[i] in self.manager.days):  # and self.period[0] != self.manager.period[0]
                        self.table.setItem(row_id, i, QTableWidgetItem(str(week[i])))
                        self.table.item(row_id, i).setBackground(QtGui.QColor(60, 170, 60))

                    elif week[i] in self.manager.days:
                        self.table.setItem(row_id, i, QTableWidgetItem(str(week[i])))
                        self.table.item(row_id, i).setBackground(QtGui.QColor(0, 0, 0, 0))

                    self.table.item(row_id, i).setFont(QtGui.QFont('Calibri', 14 + self.size))
                    self.table.item(row_id, i).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
                    self.table.item(row_id, i).setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
                row_id += 1
        else:
            for week in calendar.monthcalendar(month=self.period[0], year=self.period[1]):
                for i in range(len(week)):
                    if not week[i]:
                        self.table.setItem(row_id, i, QTableWidgetItem(str('✖')))
                    elif self.manager.user.happy_days.get(str(self.period[0])) and (week[i] in self.manager.user.happy_days.get(str(self.period[0]))):
                        self.table.setItem(row_id, i, QTableWidgetItem(str(week[i])))
                        self.table.item(row_id, i).setBackground(QtGui.QColor(181, 71, 71))
                    else:
                        self.table.setItem(row_id, i, QTableWidgetItem(str(week[i])))
                        self.table.item(row_id, i).setBackground(QtGui.QColor(0, 0, 0, 0))

                    self.table.item(row_id, i).setFont(QtGui.QFont('Calibri', 14 + self.size))
                    self.table.item(row_id, i).setTextAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
                    self.table.item(row_id, i).setFlags(QtCore.Qt.ItemFlag.ItemIsEnabled)
                row_id += 1

        self.table.horizontalHeader().setSectionResizeMode(QtWidgets.QHeaderView.ResizeMode.Stretch)
        self.table.verticalHeader().setSectionResizeMode(QtWidgets.QHeaderView.ResizeMode.Stretch)
        self.table.setHorizontalHeaderLabels([self.tr('ПН'), self.tr("ВТ"), self.tr('СР'), self.tr('ЧТ'), self.tr('ПТ'),
                                              self.tr('СБ'), self.tr('ВС'), ])

        self.table.cellClicked.connect(self.click_cell)
        self.table.cellDoubleClicked.connect(self.click_double_cell)
        self.table_layout.addWidget(self.table)
        self.vertical_layout.addLayout(self.table_layout)

    def click_double_cell(self, row, column):
        if self.table.item(row, column).text().isnumeric():
            if not (self.table.item(row, column).background() == QtGui.QColor(181, 71, 71)):
                if not self.period[0] != self.manager.period[0]:
                    self.manager.off_day(int(self.table.item(row, column).text()), is_happy_day=True)
                else:
                    self.manager.off_day(int(self.table.item(row, column).text()), is_happy_day=True, month=self.period[0])
                self.table.item(row, column).setBackground(QtGui.QColor(181, 71, 71))
                self.parent.F6.init_table_absence()
            else:
                if not self.period[0] != self.manager.period[0]:
                    self.manager.on_day(int(self.table.item(row, column).text()), is_happy_day=True)
                else:
                    self.manager.on_day(int(self.table.item(row, column).text()), is_happy_day=True, month=self.period[0])
                self.table.item(row, column).setBackground(QtGui.QColor(0, 0, 0, 0))
                self.parent.F6.init_table_absence()


class YearCalendarView(MonthCalendarView):
    def click_cell(self, row, column):
        self.click_double_cell(row, column)



class UpdateManager(QtWidgets.QDialog):
    URL_PROJECT = 'https://api.github.com/repos/IZBYSHNIK/F6/releases/latest'

    def __init__(self, parent=None):
        super(UpdateManager, self).__init__(parent)
        self.setObjectName("UpdateManager")
        self.setStyleSheet('a {font-size:18px; color: black; text-decoration: none;}')

        self.setModal(True)
        self.vertical_layout = QtWidgets.QVBoxLayout(self)

        self.title = QtWidgets.QLabel()
        self.title.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.title.setStyleSheet('font-size:18px; text-transform: uppercase;')

        self.message = QtWidgets.QLabel()
        self.message.setStyleSheet('font-size:14px;')

        self.buttons_layout = QtWidgets.QHBoxLayout()
        self.show_link = QtWidgets.QLabel()
        self.show_link.setStyleSheet('font-size:16px')
        self.show_link.setOpenExternalLinks(True)
        self.show_link.setFixedHeight(50)

        self.show_link.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)

        self.buttons_layout.addWidget(self.show_link)
        self.load_link = QtWidgets.QLabel()
        self.load_link.setStyleSheet('font-size:16px')
        self.load_link.setOpenExternalLinks(True)
        self.load_link.setFixedHeight(50)

        self.load_link.setAlignment(QtCore.Qt.AlignmentFlag.AlignCenter)
        self.buttons_layout.addWidget(self.load_link)

        self.vertical_layout.addWidget(self.title)
        self.vertical_layout.addWidget(self.message)
        self.vertical_layout.addLayout(self.buttons_layout)

        self.retranslateUi()

    def retranslateUi(self):
        self.setWindowTitle(self.tr('Обновление'))

    def fill_data(self, data):
        self.title.setText(data['title'])
        self.message.setText(data['message'])
        self.show_link.setText(f'<a href="{data["show_url"]}">' + self.tr('Посмотреть') + '</a>')
        self.load_link.setText(f'<a href="{data["download_url"]}">' + self.tr('Скачать') + '</a>')

    def link_host(self):
        data = {}
        try:
            update_massage = json.loads(requests.api.get(self.URL_PROJECT,
                                                         params={'User-Agent:': str(platform.node())}).text)
        except requests.exceptions.ConnectionError as f:
            return data

        if not update_massage:
            return {}

        version = [int(i) if i.isnumeric() else i for i in
                   update_massage['tag_name'][1:].split('.')]

        data = {
            'version': version,
            'title': update_massage['name'],
            'message': update_massage['body'],
            'show_url': update_massage['html_url'],
            'download_url': update_massage['assets'][0]['browser_download_url'],
        }

        return data

    @staticmethod
    def check_new_version(old, new):
        if old < new:
            return True
        return False

    def show(self):
        data = self.link_host()

        if not data:
            return
        if not self.check_new_version([int(i) if i.isnumeric() else i for i in VERSION.split('.')], data['version']):
            return

        self.fill_data(data)

        super().show()


class SettingsWindows(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super(SettingsWindows, self).__init__(parent)
        self.setObjectName("SettingsWindows")
        self.setFixedSize(400, 200)
        self.setModal(True)
        self.verticalLayout = QtWidgets.QVBoxLayout(self)
        self.verticalLayout.setObjectName("verticalLayout")
        self.scrollArea = QtWidgets.QScrollArea(self)
        self.scrollArea.setWidgetResizable(True)
        self.scrollArea.setObjectName("scrollArea")
        self.scrollAreaWidgetContents = QtWidgets.QWidget()
        self.scrollAreaWidgetContents.setGeometry(QtCore.QRect(0, 0, 376, 120))
        self.scrollAreaWidgetContents.setObjectName("scrollAreaWidgetContents")
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self.scrollAreaWidgetContents)
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.title_messages = QtWidgets.QLabel(self.scrollAreaWidgetContents)
        self.title_messages.setObjectName("label")
        self.verticalLayout_2.addWidget(self.title_messages)
        self.lineEdit = QtWidgets.QSpinBox(self.scrollAreaWidgetContents)
        self.lineEdit.setObjectName("lineEdit")
        self.verticalLayout_2.addWidget(self.lineEdit)
        self.messages_error = QtWidgets.QTextBrowser(self.scrollAreaWidgetContents)
        self.messages_error.setStyleSheet("color: red; background: none; border: none;")
        self.messages_error.setFixedHeight(50)
        self.verticalLayout_2.addWidget(self.messages_error)
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                           QtWidgets.QSizePolicy.Policy.Expanding)
        self.verticalLayout_2.addItem(spacerItem)
        self.scrollArea.setWidget(self.scrollAreaWidgetContents)
        self.verticalLayout.addWidget(self.scrollArea)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setContentsMargins(-1, 20, -1, -1)
        self.horizontalLayout.setObjectName("horizontalLayout")
        spacerItem1 = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Policy.Expanding,
                                            QtWidgets.QSizePolicy.Policy.Minimum)
        self.horizontalLayout.addItem(spacerItem1)
        self.save_button = QtWidgets.QPushButton(self)
        self.save_button.setObjectName("save_button")
        self.horizontalLayout.addWidget(self.save_button)
        self.cancel_button = QtWidgets.QPushButton(self)
        self.cancel_button.setObjectName("cancel_button")
        self.horizontalLayout.addWidget(self.cancel_button)

        self.verticalLayout.addLayout(self.horizontalLayout)

        self.result = None

        self.retranslateUi()
        QtCore.QMetaObject.connectSlotsByName(self)

        self.cancel_button.clicked.connect(self.clicked_cancel)
        self.save_button.clicked.connect(self.clicked_save)

    def retranslateUi(self):
        self.setWindowTitle(self.tr("Рабочий день"))
        self.title_messages.setText(self.tr(""))
        self.cancel_button.setText(self.tr("Отмена"))
        self.save_button.setText(self.tr("Сохранить"))

    def clicked_cancel(self):
        self.close()

    def check_value(self):
        if self.lineEdit.text().isnumeric() and 31 >= int(self.lineEdit.text()) > 0:
            return int(self.lineEdit.text())
        raise ValueError(self.tr('Введен недопустимый номер дня'))

    def clicked_save(self):
        try:
            res = self.check_value()
        except ValueError as f:
            self.messages_error.setText(str(f))
        else:
            self.result = res
            self.close()


class SettingsData(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super(SettingsData, self).__init__(parent)
        self.setObjectName("SettingsData")
        self.resize(431, 160)
        self.status = 'finished'
        self.setFixedSize(QtCore.QSize(430, 160))
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.verticalLayout = QtWidgets.QVBoxLayout()
        self.verticalLayout.setObjectName("verticalLayout")
        self.dateEdit = QtWidgets.QDateEdit(self)
        self.dateEdit.setDateTime(QtCore.QDateTime(QtCore.QDate.currentDate(), QtCore.QTime(0, 0, 0)))
        self.dateEdit.setObjectName("dateEdit")
        self.dateEdit.setDisplayFormat("MM.yyyy")
        self.verticalLayout.addWidget(self.dateEdit)
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                           QtWidgets.QSizePolicy.Policy.Expanding)
        self.verticalLayout.addItem(spacerItem)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.create_new_table_pushButton = QtWidgets.QPushButton(self)
        self.create_new_table_pushButton.setObjectName("create_new_table_pushButton")
        self.horizontalLayout.addWidget(self.create_new_table_pushButton)

        self.cancel_pushButton = QtWidgets.QPushButton(self)
        self.cancel_pushButton.setMinimumSize(QtCore.QSize(0, 0))
        self.cancel_pushButton.setMaximumSize(QtCore.QSize(16777215, 16777215))
        self.cancel_pushButton.setObjectName("cancel_pushButton")
        self.horizontalLayout.addWidget(self.cancel_pushButton)
        self.verticalLayout.addLayout(self.horizontalLayout)
        self.verticalLayout_2.addLayout(self.verticalLayout)

        self.retranslateUi()
        QtCore.QMetaObject.connectSlotsByName(self)

        self.create_new_table_pushButton.clicked.connect(self.clicked_create_table)
        self.cancel_pushButton.clicked.connect(self.clicked_cancel)

    def retranslateUi(self):
        self.setWindowTitle(self.tr("Период"))
        self.create_new_table_pushButton.setText(self.tr("Изменить"))
        self.cancel_pushButton.setText(self.tr("Отмена"))

    def clicked_create_table(self):
        self.close()

    def clicked_cancel(self):
        self.status = 'exit'
        self.close()


class Push(QtWidgets.QPushButton):
    def __init__(self, parent, base_weight, base_height, growth, tool_tip=None, icon_path=None):
        super().__init__()
        self.setStyleSheet('border: none;')
        self.base_height = base_height
        self.base_weight = base_weight
        self.growth = growth
        self.setIconSize(QtCore.QSize(self.base_weight, self.base_height))
        if tool_tip:
            self.setToolTip(tool_tip)
        if icon_path:
            self.setIcon(QtGui.QIcon(icon_path))

    def enterEvent(self, e):
        self.setIconSize(QtCore.QSize(self.base_weight + self.growth, self.base_height + self.growth))

    def leaveEvent(self, e):
        self.setIconSize(QtCore.QSize(self.base_weight, self.base_height))


class WindowSetPassword(QtWidgets.QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.parent = parent
        self.setObjectName("Form")
        self.resize(500, 300)
        self.setMaximumSize(QtCore.QSize(500, 300))
        self.verticalLayout_2 = QtWidgets.QVBoxLayout(self)
        self.verticalLayout_2.setObjectName("verticalLayout_2")
        self.verticalLayout = QtWidgets.QVBoxLayout()
        self.verticalLayout.setObjectName("verticalLayout")
        self.label = QtWidgets.QLabel(self)
        self.label.setObjectName("label")
        self.verticalLayout.addWidget(self.label)
        self.old_password_edit = QtWidgets.QLineEdit(self)
        self.old_password_edit.setEnabled(True)
        self.old_password_edit.setObjectName("old_password_edit")
        self.verticalLayout.addWidget(self.old_password_edit)
        self.label_2 = QtWidgets.QLabel(self)
        self.label_2.setObjectName("label_2")
        self.verticalLayout.addWidget(self.label_2)
        self.new_password_edit1 = QtWidgets.QLineEdit(self)
        self.new_password_edit1.setObjectName("new_password_edit1")
        self.verticalLayout.addWidget(self.new_password_edit1)
        self.label_3 = QtWidgets.QLabel(self)
        self.label_3.setObjectName("label_3")
        self.verticalLayout.addWidget(self.label_3)
        self.new_password_edit2 = QtWidgets.QLineEdit(self)
        self.new_password_edit2.setObjectName("new_password_edit2")
        self.verticalLayout.addWidget(self.new_password_edit2)
        self.fielderrors = QtWidgets.QTextBrowser(self)
        self.fielderrors.setObjectName("fielderrors")
        self.verticalLayout.addWidget(self.fielderrors)
        spacerItem = QtWidgets.QSpacerItem(20, 40, QtWidgets.QSizePolicy.Policy.Minimum,
                                           QtWidgets.QSizePolicy.Policy.Expanding)
        self.verticalLayout.addItem(spacerItem)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        self.save_password_pushbutton = QtWidgets.QPushButton(self)
        self.save_password_pushbutton.setObjectName("save_password_pushbutton")
        self.horizontalLayout.addWidget(self.save_password_pushbutton)
        self.line = QtWidgets.QFrame(self)
        self.line.setFrameShape(QtWidgets.QFrame.Shape.VLine)
        self.line.setFrameShadow(QtWidgets.QFrame.Shadow.Sunken)
        self.line.setObjectName("line")
        self.horizontalLayout.addWidget(self.line)
        self.cancel_pushbutton = QtWidgets.QPushButton(self)
        self.cancel_pushbutton.setObjectName("cancel_pushbutton")
        self.horizontalLayout.addWidget(self.cancel_pushbutton)
        self.verticalLayout.addLayout(self.horizontalLayout)
        self.verticalLayout_2.addLayout(self.verticalLayout)
        self.retranslateUi()

        self.add_function()

    def retranslateUi(self, ):
        self.setWindowTitle(self.tr("Пароль"))
        self.label.setText(self.tr("Старый пароль"))
        self.label_2.setText(self.tr("Новый пароль"))
        self.label_3.setText(self.tr("Повторите новый пароль"))
        self.save_password_pushbutton.setText(self.tr("Сохранить"))
        self.cancel_pushbutton.setText(self.tr("Отмена"))

        self.label.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.label_2.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.label_3.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.save_password_pushbutton.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.cancel_pushbutton.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.fielderrors.setFont(QtGui.QFont(CM.NAME_FONT, 14 + CM.ADD_FONT_SIZE))
        self.fielderrors.setStyleSheet('color: red')

    def add_function(self):
        self.cancel_pushbutton.clicked.connect(self.click_cancel)
        self.save_password_pushbutton.clicked.connect(self.click_save)

    def click_cancel(self):
        self.close()

    def checking_parametrs(self):
        if not self.old_password_edit.text() == USER_MANAGER.user.password:
            raise ValueError(self.tr('Введен неверный текущий пароль'))
        if not self.new_password_edit2.text() == self.new_password_edit1.text():
            raise ValueError(self.tr('Пароли не соответствуют'))

    def click_save(self):
        self.fielderrors.clear()
        try:
            self.checking_parametrs()
        except ValueError as m:
            self.fielderrors.setText(str(m))

        else:
            try:
                USER_MANAGER.user.password = self.new_password_edit1.text()
            except ValueError as m:
                self.fielderrors.setText(str(m))
            else:
                USER_MANAGER.user.save_user()
                MANAGER_STUDENTS.user = USER_MANAGER.user

                MANAGER_STUDENTS.save_students()
                self.parent.get_down_message(self.tr('Пароль успешно изменен'))
                self.close()



class ControlerWindows(QtWidgets.QWidget):
    def __init__(self, splash_screen: SplashScreen, auth: Auth, regist: Regist, main: MainWindow):
        self.splash_screen = splash_screen()
        self.splash_screen.show()
        self.auth = auth()
        self.load_splash_screen(0, 31)
        self.regist = regist()
        self.load_splash_screen(31, 61)
        self.main = main()
        self.load_splash_screen(61, 101)

        self.auth.closeEvent = self.close_event_by_auth
        self.main.closeEvent = self.close_event_by_main
        self.regist.closeEvent = self.close_event_by_regist
        self.splash_screen.close()

    def load_splash_screen(self, starts=0, end=0):
        for i in range(starts, end):
            self.splash_screen.progressBar.setValue(i)
            QtCore.QThread.msleep(4)

    def close_event_by_auth(self, e):
        if self.auth.status == 1:
            self.main.init_students_manager('students.json')
            self.main.show()
        elif self.auth.status == 3:
            self.regist.show()
        else:
            self.auth.close()
        self.auth.status = 0
        self.auth.closeEvent = self.close_event_by_auth

    def close_event_by_main(self, e):
        if self.main.status == 2:
            self.auth.update_users()
            self.auth.show()
        elif self.main.status == 3:
            self.regist.show()
        elif self.main.status == 4:
            self.is_exit(self.main)
            self.main = type(self.main)()
            self.auth.show()
            CM.IS_CHANGE = False

        else:
            self.main.close()
            self.is_exit(self.main)

        self.main.status = 0
        self.main.closeEvent = self.close_event_by_main

    def close_event_by_regist(self, e):
        if self.regist.status == 1:
            self.auth.update_users()
            self.main.init_students_manager('students.json')
            self.main.show()
        elif self.regist.status == 2:
            self.auth.show()
        else:
            self.regist.close()
        self.regist.status = 0

    def show(self):
        if USER_MANAGER.users_id:
            self.auth.show()
        else:
            self.regist.show()

    def is_exit(self, parent):
        if CM.IS_CHANGE:
            message = QtWidgets.QMessageBox.question(self.main, self.tr('Сохранение изменений'),
                                                     self.tr("Сохранить изменения?"),
                                                     QtWidgets.QMessageBox.StandardButton.Yes | QtWidgets.QMessageBox.StandardButton.No)
            if message == message.Yes:
                MANAGER_STUDENTS.save_students()

    def retranslateUi(self):
        self.main.retranslateUi()
        self.auth.retranslateUi()
        self.regist.retranslateUi()
        self.splash_screen.retranslateUi()


def set_language(*args, **kwargs):
    global USER_MANAGER
    CM.CURRENT_LANGUAGE = next(LANGUAGES)
    if translator.load(os.path.join(CONTENT_PATH, 'languages', CM.CURRENT_LANGUAGE)):
        app.installTranslator(translator)
        USER_MANAGER.parametrs['language'] = CM.CURRENT_LANGUAGE
    else:
        translator.load("qt_" + 'ru_RU', QLibraryInfo.location(QLibraryInfo.TranslationsPath))
        app.installTranslator(translator)
        USER_MANAGER.parametrs['language'] = 'russia'
    USER_MANAGER.save_users()
    windows.retranslateUi()


try:
    filename = os.path.join(CONTENT_PATH, 'sounds', "logo.mp3")
    player = QMediaPlayer()
    audio_output = QAudioOutput()
    player.setAudioOutput(audio_output)
    player.setSource(QUrl.fromLocalFile(filename))
    audio_output.setVolume(0.15)
    player.play()
except:
    print('No sound')


app = QtWidgets.QApplication(sys.argv)
translator = QTranslator(app)

if translator.load(os.path.join(BASE_PATH, 'languages', str(USER_MANAGER.parametrs.get('language')))):
    app.installTranslator(translator)
    CM.CURRENT_LANGUAGE = USER_MANAGER.parametrs.get('language')
else:
    translator.load("qt_" + 'ru_RU', QLibraryInfo.location(QLibraryInfo.TranslationsPath))
    app.installTranslator(translator)
    CM.CURRENT_LANGUAGE = 'russia'

USER_MANAGER.save_users()





windows = ControlerWindows(SplashScreen, Auth, Regist, MainWindow)
windows.show()

status = app.exec()
if USER_MANAGER.user:
    USER_MANAGER.user.save_user()

print(status, '-')
sys.exit(status)
